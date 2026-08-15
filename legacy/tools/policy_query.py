import json
import os
import re
import glob
from typing import Dict, List, Optional



class PolicyLibrary:
    """政策文件库（意见/通知/办法/规定等）

    私有成员以 _ 开头。
    支持 PDF 和 docx 两种格式。
    """

    # 政策类型关键词
    _POLICY_TYPES = {
        "意见": ["意见"],
        "通知": ["通知"],
        "办法": ["办法"],
        "规定": ["规定", "规则"],
        "决定": ["决定"],
        "公告": ["公告"],
        "批复": ["批复"],
        "函": ["函"],
    }

    # 文号正则：如"国发〔2024〕1号"、"X政发〔2023〕15号"
    _doc_number_pattern = re.compile(
        r'[（）\(\)\u201c\u201d"\'][^（）\(\)\u201c\u201d"\'的]{0,20}'
        r'[〔\[][\d]{4}[〕\]]\s*第?\s*\d+\s*号?'
    )
    # 简化版文号
    _doc_number_simple = re.compile(
        r'[一-龥]{1,5}〔\d{4}〕\d+号'
    )

    def __init__(self, policy_dir: str = "./policies"):
        self._policy_dir = policy_dir
        self._cache = {}  # {文件名: {元数据+段落}}
        self._index = {}  # {关键词: [文件名]}
        self._build_index()

    # ---------- 私有方法 ----------
    def _extract_text(self, file_path: str) -> str:
        """根据扩展名选择解析器"""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self._extract_pdf(file_path)
        elif ext == ".docx":
            return self._extract_docx(file_path)
        return ""

    def _extract_pdf(self, path: str) -> str:
        try:
            import PyPDF2
        except ImportError:
            raise ImportError("请先安装 PyPDF2: pip install PyPDF2")
        texts = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                t = page.extract_text() or ""
                if t.strip():
                    texts.append(t.strip())
        return "\n".join(texts)

    def _extract_docx(self, path: str) -> str:
        from docx import Document
        doc = Document(path)
        texts = []
        for para in doc.paragraphs:
            t = para.text.strip()
            if t:
                texts.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        texts.append(t)
        return "\n".join(texts)

    def _detect_policy_type(self, title: str) -> str:
        """从标题判断政策类型"""
        for ptype, keywords in self._POLICY_TYPES.items():
            for kw in keywords:
                if kw in title:
                    return ptype
        return "其他"

    def _extract_doc_number(self, text: str) -> str:
        """提取文号"""
        m = self._doc_number_simple.search(text[:2000])
        if m:
            return m.group(0)
        m = self._doc_number_pattern.search(text[:2000])
        if m:
            return m.group(0)
        return ""

    def _extract_issuer(self, text: str) -> str:
        """提取发文机关（通常在标题下方或正文开头）"""
        # 常见模式："XX部 XX局"、"XX省人民政府"
        patterns = [
            r'^([\u4e00-\u9fa5]{2,10}(?:部|委|办|局|署|院|厅|人民政府|办公厅))',
            r'([\u4e00-\u9fa5]{2,8}(?:委员会|办公室|领导小组))',
        ]
        lines = text[:1000].split("\n")
        for line in lines:
            for pat in patterns:
                m = re.search(pat, line)
                if m:
                    return m.group(1)
        return ""

    def _split_sections(self, text: str) -> List[Dict]:
        """将正文按段落/条目切分，便于引用"""
        lines = [l.strip() for l in text.split("\n") if l.strip()]
        sections = []
        current_title = ""
        current_content = []

        # 匹配"一、" "（一）" "1." "第一条" 等
        heading_patterns = [
            re.compile(r'^[一二三四五六七八九十]+、'),
            re.compile(r'^（[一二三四五六七八九十]+）'),
            re.compile(r'^\d+[\.、．]'),
            re.compile(r'^第[一二三四五六七八九十百零\d]+条'),
        ]

        for line in lines:
            is_heading = any(p.match(line) for p in heading_patterns)
            if is_heading:
                if current_title or current_content:
                    sections.append({
                        "title": current_title or f"段落{len(sections) + 1}",
                        "content": "\n".join(current_content),
                    })
                current_title = line
                current_content = []
            else:
                # 跳过纯元数据行（日期、文号行等）
                if len(line) < 5:
                    continue
                current_content.append(line)

        if current_title or current_content:
            sections.append({
                "title": current_title or f"段落{len(sections) + 1}",
                "content": "\n".join(current_content),
            })

        return sections

    def _build_index(self):
        if not os.path.exists(self._policy_dir):
            os.makedirs(self._policy_dir, exist_ok=True)
            return

        for ext in ("*.pdf", "*.docx", "*.PDF", "*.DOCX"):
            for fp in glob.glob(os.path.join(self._policy_dir, ext)):
                fname = os.path.basename(fp)
                try:
                    text = self._extract_text(fp)
                    sections = self._split_sections(text)

                    # 从文件名推测标题
                    title = re.sub(r'\.(pdf|docx)$', '', fname, flags=re.I)
                    title = re.sub(r'【.*?】', '', title).strip()

                    # 从正文第一行尝试获取更完整的标题
                    lines = text.split("\n")
                    for line in lines[:5]:
                        line = line.strip()
                        if 5 < len(line) < 100 and not line.startswith(" "):
                            if any(kw in line for kw in ["意见", "通知", "办法", "规定"]):
                                title = line
                                break

                    doc_num = self._extract_doc_number(text)
                    issuer = self._extract_issuer(text)
                    ptype = self._detect_policy_type(title)

                    self._cache[fname] = {
                        "title": title,
                        "file_path": fp,
                        "type": ptype,
                        "issuer": issuer,
                        "doc_number": doc_num,
                        "total_sections": len(sections),
                        "sections": sections,
                        "full_text": text,
                        "text_length": len(text),
                    }

                    # 建索引
                    for i in range(len(title) - 1):
                        self._index.setdefault(title[i:i + 2], []).append(fname)
                    self._index.setdefault(title, []).append(fname)
                    if doc_num:
                        self._index.setdefault(doc_num, []).append(fname)

                except Exception as e:
                    print(f"⚠️ 解析 {fname} 失败: {e}")

    def _find_policy(self, keyword: str) -> Optional[str]:
        if keyword in self._index:
            return self._index[keyword][0]
        for fname, info in self._cache.items():
            if keyword in info["title"] or keyword in info.get("doc_number", ""):
                return fname
        # bigram 模糊匹配
        best = None
        best_score = 0
        for i in range(len(keyword) - 1):
            bg = keyword[i:i + 2]
            if bg in self._index:
                for fn in self._index[bg]:
                    score = sum(1 for j in range(len(keyword) - 1)
                                if keyword[j:j + 2] in self._cache[fn]["title"])
                    if score > best_score:
                        best_score = score
                        best = fn
        return best

    # ---------- 公开方法 ----------
    def list_policies(self, policy_type: str = "") -> List[Dict]:
        """列出所有政策文件，可按类型过滤"""
        results = []
        for info in self._cache.values():
            if policy_type and info["type"] != policy_type:
                continue
            results.append({
                "title": info["title"],
                "type": info["type"],
                "issuer": info["issuer"],
                "doc_number": info["doc_number"],
                "sections": info["total_sections"],
            })
        return results

    def get_section(self, keyword: str, section_ref: str = "") -> Dict:
        """获取指定政策文件的段落内容

        Args:
            keyword: 文件标题关键词
            section_ref: 段落标识，如"一、" "（二）" "第一条"。留空返回目录
        """
        fn = self._find_policy(keyword)
        if not fn:
            return {"found": False, "message": f"未找到包含'{keyword}'的政策文件"}

        info = self._cache[fn]
        sections = info["sections"]

        if not section_ref:
            return {
                "found": True,
                "title": info["title"],
                "type": info["type"],
                "issuer": info["issuer"],
                "doc_number": info["doc_number"],
                "total_sections": len(sections),
                "toc": [s["title"] for s in sections],
            }

        # 精确匹配段落
        for s in sections:
            if section_ref in s["title"] or s["title"].startswith(section_ref):
                return {
                    "found": True,
                    "title": info["title"],
                    "section": s["title"],
                    "content": s["content"],
                }

        # 模糊匹配
        for s in sections:
            if section_ref in s["content"][:50]:
                return {
                    "found": True,
                    "title": info["title"],
                    "section": s["title"],
                    "content": s["content"],
                }

        return {
            "found": True,
            "title": info["title"],
            "section_requested": section_ref,
            "message": f"未找到该段落，共{len(sections)}个段落",
        }

    def search_policies(self, keyword: str) -> Dict:
        """在所有政策文件中检索关键词"""
        results = []
        for info in self._cache.values():
            # 在标题和全文中搜索
            if keyword in info["title"] or keyword in info["full_text"]:
                # 找到具体段落
                matched_sections = []
                for s in info["sections"]:
                    if keyword in s["title"] or keyword in s["content"]:
                        matched_sections.append({
                            "section": s["title"],
                            "content": s["content"][:500],
                        })
                results.append({
                    "title": info["title"],
                    "type": info["type"],
                    "issuer": info["issuer"],
                    "doc_number": info["doc_number"],
                    "matches": matched_sections[:5],
                })
        return {"keyword": keyword, "total": len(results), "results": results[:10]}

    def get_full_text(self, keyword: str, max_chars: int = 8000) -> str:
        """获取完整文本（截断），用于 LLM context"""
        fn = self._find_policy(keyword)
        if not fn:
            return ""
        text = self._cache[fn]["full_text"]
        info = self._cache[fn]
        header = f"《{info['title']}》（{info['issuer']}，{info['doc_number']}）\n\n"
        return header + text[:max_chars]

    def reload(self):
        self._cache.clear()
        self._index.clear()
        self._build_index()


# ---------- 统一工具函数 ----------
def search_policy_query(action: str, keyword: str = "", section_ref: str = "",
                 policy_type: str = "", policy_dir: str = "./policies") -> str:
    """
    统一政策文件查询工具。

    参数：
        action: "list" | "get_section" | "search" | "full_text"
        keyword: 文件标题关键词（get_section/search/full_text 时需要）
        section_ref: 段落标识，如"一、" "（二）"（get_section 时可选）
        policy_type: 政策类型过滤（list 时可选）：意见/通知/办法/规定/决定
        policy_dir: 政策文件所在目录

    返回：JSON 字符串
    """
    lib = PolicyLibrary(policy_dir=policy_dir)

    if action == "list":
        results = lib.list_policies(policy_type)
        return json.dumps({"total": len(results), "policies": results}, ensure_ascii=False)

    elif action == "get_section":
        result = lib.get_section(keyword, section_ref)
        return json.dumps(result, ensure_ascii=False)

    elif action == "search":
        result = lib.search_policies(keyword)
        return json.dumps(result, ensure_ascii=False)

    elif action == "full_text":
        text = lib.get_full_text(keyword)
        return json.dumps({"text": text}, ensure_ascii=False)

    else:
        return json.dumps({"error": f"未知 action: {action}"}, ensure_ascii=False)

