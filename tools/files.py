"""案件电子卷宗的文件上传与后续处理。

覆盖上传落盘与哈希、版本链、解析与页级 OCR、质量状态、分块、脱敏、
外发门控、人工修正，以及提供给 Agent 调用的材料工具。
"""

from __future__ import annotations

import hashlib
import io
import json
import mimetypes
import os
import re
import sqlite3
import uuid
from contextlib import contextmanager
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterator, Optional, Protocol

from app.config import (
    ALLOWED_MATERIAL_EXTENSIONS,
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    DATABASE_PATH,
    MATERIAL_STORAGE_DIR,
    MAX_UPLOAD_BYTES,
    OCR_LOW_CONFIDENCE_THRESHOLD,
    OCR_MAX_PAGE_RETRIES,
    OCR_TEXT_DENSITY_THRESHOLD,
    PARSER_VERSION,
    REDACTION_STORAGE_DIR,
)

# ---------- 状态与错误码 ----------

MATERIAL_STATUSES = {
    "UPLOADED",
    "PARSING",
    "PARSED",
    "NEEDS_OCR_REVIEW",
    "OCR_FAILED",
    "DUPLICATE_PENDING",
    "FAILED",
    "DELETED",
}

ERROR_CODES = {
    "UNSUPPORTED_TYPE": "MATERIAL_UNSUPPORTED_TYPE",
    "FILE_TOO_LARGE": "MATERIAL_FILE_TOO_LARGE",
    "CORRUPT_FILE": "MATERIAL_CORRUPT_FILE",
    "ENCRYPTED_FILE": "MATERIAL_ENCRYPTED_FILE",
    "OCR_FAILED": "MATERIAL_OCR_FAILED",
    "PARSE_FAILED": "MATERIAL_PARSE_FAILED",
    "AUTH_DENIED": "MATERIAL_AUTH_DENIED",
    "EGRESS_DENIED": "MATERIAL_EGRESS_DENIED",
    "NOT_FOUND": "MATERIAL_NOT_FOUND",
    "DUPLICATE_PENDING": "MATERIAL_DUPLICATE_PENDING",
}


class MaterialError(Exception):
    def __init__(self, code: str, message: str, details: dict | None = None):
        self.code = code
        self.message = message
        self.details = details or {}
        super().__init__(f"{code}: {message}")

    def to_dict(self) -> dict[str, Any]:
        return {"error_code": self.code, "message": self.message, "details": self.details}


# ---------- 数据库 ----------

SCHEMA_SQL = """
CREATE TABLE IF NOT EXISTS users (
    id VARCHAR(64) PRIMARY KEY,
    display_name VARCHAR(100) NOT NULL,
    role VARCHAR(32) NOT NULL
);

CREATE TABLE IF NOT EXISTS case_pools (
    id VARCHAR(36) PRIMARY KEY,
    name VARCHAR(120) NOT NULL,
    purpose TEXT NOT NULL,
    owner_user_id VARCHAR(64) NOT NULL,
    created_at DATETIME NOT NULL
);

CREATE TABLE IF NOT EXISTS case_pool_members (
    id VARCHAR(36) PRIMARY KEY,
    case_pool_id VARCHAR(36) NOT NULL,
    user_id VARCHAR(64) NOT NULL,
    member_role VARCHAR(32) NOT NULL,
    UNIQUE(case_pool_id, user_id)
);

CREATE TABLE IF NOT EXISTS cases (
    id VARCHAR(36) PRIMARY KEY,
    case_pool_id VARCHAR(36) NOT NULL,
    name VARCHAR(160) NOT NULL,
    case_number VARCHAR(100) NOT NULL,
    created_by VARCHAR(64) NOT NULL,
    created_at DATETIME NOT NULL,
    UNIQUE(case_pool_id, case_number)
);

CREATE TABLE IF NOT EXISTS documents (
    id VARCHAR(36) PRIMARY KEY,
    case_id VARCHAR(36) NOT NULL,
    filename VARCHAR(255) NOT NULL,
    stored_name VARCHAR(255) NOT NULL,
    content_type VARCHAR(120) NOT NULL,
    size BIGINT NOT NULL,
    sha256 VARCHAR(64) NOT NULL,
    uploaded_by VARCHAR(64) NOT NULL,
    created_at DATETIME NOT NULL,
    status VARCHAR(32) NOT NULL DEFAULT 'UPLOADED',
    current_version_id VARCHAR(36),
    deleted_at DATETIME,
    quality_summary_json TEXT NOT NULL DEFAULT '{}'
);

CREATE TABLE IF NOT EXISTS document_versions (
    id VARCHAR(36) PRIMARY KEY,
    document_id VARCHAR(36) NOT NULL,
    version_no INTEGER NOT NULL,
    parent_version_id VARCHAR(36),
    source_type VARCHAR(32) NOT NULL,
    sha256 VARCHAR(64) NOT NULL,
    storage_path TEXT NOT NULL,
    content_type VARCHAR(120) NOT NULL,
    size BIGINT NOT NULL,
    parser_version VARCHAR(64) NOT NULL,
    status VARCHAR(32) NOT NULL,
    is_current INTEGER NOT NULL DEFAULT 0,
    is_active INTEGER NOT NULL DEFAULT 1,
    quality_summary_json TEXT NOT NULL DEFAULT '{}',
    error_code VARCHAR(64),
    error_message TEXT,
    created_by VARCHAR(64) NOT NULL,
    created_at DATETIME NOT NULL,
    UNIQUE(document_id, version_no)
);

CREATE TABLE IF NOT EXISTS parse_jobs (
    id VARCHAR(36) PRIMARY KEY,
    document_version_id VARCHAR(36) NOT NULL,
    status VARCHAR(32) NOT NULL,
    attempt INTEGER NOT NULL DEFAULT 0,
    max_attempts INTEGER NOT NULL DEFAULT 2,
    error_code VARCHAR(64),
    error_message TEXT,
    started_at DATETIME,
    finished_at DATETIME,
    created_at DATETIME NOT NULL
);

CREATE TABLE IF NOT EXISTS document_pages (
    id VARCHAR(36) PRIMARY KEY,
    document_version_id VARCHAR(36) NOT NULL,
    page_no INTEGER NOT NULL,
    source VARCHAR(32) NOT NULL,
    text TEXT NOT NULL DEFAULT '',
    text_density REAL NOT NULL DEFAULT 0,
    avg_confidence REAL,
    min_confidence REAL,
    bbox_json TEXT NOT NULL DEFAULT '[]',
    lines_json TEXT NOT NULL DEFAULT '[]',
    quality_flags_json TEXT NOT NULL DEFAULT '[]',
    status VARCHAR(32) NOT NULL,
    retry_count INTEGER NOT NULL DEFAULT 0,
    error_code VARCHAR(64),
    error_message TEXT,
    UNIQUE(document_version_id, page_no)
);

CREATE TABLE IF NOT EXISTS document_chunks (
    id VARCHAR(64) PRIMARY KEY,
    document_version_id VARCHAR(36) NOT NULL,
    ordinal INTEGER NOT NULL,
    page_start INTEGER NOT NULL,
    page_end INTEGER NOT NULL,
    char_start INTEGER NOT NULL,
    char_end INTEGER NOT NULL,
    bbox_json TEXT NOT NULL DEFAULT '[]',
    text_raw TEXT NOT NULL,
    text_redacted TEXT NOT NULL,
    text_sha256 VARCHAR(64) NOT NULL,
    parser_version VARCHAR(64) NOT NULL,
    quality_flags_json TEXT NOT NULL DEFAULT '[]',
    is_active INTEGER NOT NULL DEFAULT 1,
    stale INTEGER NOT NULL DEFAULT 0,
    UNIQUE(document_version_id, ordinal)
);

CREATE TABLE IF NOT EXISTS redaction_items (
    id VARCHAR(36) PRIMARY KEY,
    document_version_id VARCHAR(36) NOT NULL,
    chunk_id VARCHAR(64),
    sens_type VARCHAR(32) NOT NULL,
    start_offset INTEGER NOT NULL,
    end_offset INTEGER NOT NULL,
    placeholder VARCHAR(64) NOT NULL,
    map_ref VARCHAR(128) NOT NULL,
    created_at DATETIME NOT NULL
);

CREATE TABLE IF NOT EXISTS material_audit_events (
    id VARCHAR(36) PRIMARY KEY,
    event_type VARCHAR(64) NOT NULL,
    actor_user_id VARCHAR(64),
    case_id VARCHAR(36),
    document_id VARCHAR(36),
    document_version_id VARCHAR(36),
    decision VARCHAR(32) NOT NULL,
    reason TEXT NOT NULL,
    detail_json TEXT NOT NULL DEFAULT '{}',
    created_at DATETIME NOT NULL
);

CREATE INDEX IF NOT EXISTS idx_documents_case ON documents(case_id);
CREATE INDEX IF NOT EXISTS idx_versions_document ON document_versions(document_id);
CREATE INDEX IF NOT EXISTS idx_pages_version ON document_pages(document_version_id);
CREATE INDEX IF NOT EXISTS idx_chunks_version ON document_chunks(document_version_id);
CREATE INDEX IF NOT EXISTS idx_redaction_version ON redaction_items(document_version_id);
"""


def new_id() -> str:
    return str(uuid.uuid4())


def utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def get_connection(db_path=None) -> sqlite3.Connection:
    conn = sqlite3.connect(str(db_path or DATABASE_PATH), check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def init_db(db_path=None) -> None:
    """建表并补齐旧库缺失的列，可重复执行。"""
    conn = get_connection(db_path)
    try:
        conn.executescript(SCHEMA_SQL)
        existing = {
            row["name"] for row in conn.execute("PRAGMA table_info(documents)").fetchall()
        }
        for column, ddl in (
            ("status", "status VARCHAR(32) NOT NULL DEFAULT 'UPLOADED'"),
            ("current_version_id", "current_version_id VARCHAR(36)"),
            ("deleted_at", "deleted_at DATETIME"),
            ("quality_summary_json", "quality_summary_json TEXT NOT NULL DEFAULT '{}'"),
        ):
            if column not in existing:
                conn.execute(f"ALTER TABLE documents ADD COLUMN {ddl}")
        conn.commit()
    finally:
        conn.close()


@contextmanager
def db_session(db_path=None) -> Iterator[sqlite3.Connection]:
    conn = get_connection(db_path)
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def _row(conn, sql: str, params=()) -> Optional[dict[str, Any]]:
    row = conn.execute(sql, params).fetchone()
    return dict(row) if row else None


def _rows(conn, sql: str, params=()) -> list[dict[str, Any]]:
    return [dict(row) for row in conn.execute(sql, params).fetchall()]


def _insert(conn, table: str, fields: dict[str, Any]) -> None:
    columns = list(fields)
    placeholders = ",".join("?" for _ in columns)
    conn.execute(
        f"INSERT INTO {table}({','.join(columns)}) VALUES ({placeholders})",
        [fields[column] for column in columns],
    )


def _update(conn, table: str, row_id: str, fields: dict[str, Any]) -> None:
    if not fields:
        return
    assignments = ", ".join(f"{key} = ?" for key in fields)
    conn.execute(
        f"UPDATE {table} SET {assignments} WHERE id = ?",
        [*fields.values(), row_id],
    )


def get_case(conn, case_id: str) -> Optional[dict[str, Any]]:
    return _row(conn, "SELECT * FROM cases WHERE id = ?", (case_id,))


def ensure_demo_case(conn, case_id: str | None = None) -> str:
    """为本地导入准备可用案件：指定 id 时建该案件，未指定时复用任一已有案件。"""
    if case_id and get_case(conn, case_id):
        return case_id
    if case_id is None:
        existing = _row(conn, "SELECT id FROM cases LIMIT 1")
        if existing:
            return existing["id"]

    now = utc_now()
    if not _row(conn, "SELECT id FROM users WHERE id = ?", ("system",)):
        _insert(
            conn,
            "users",
            {"id": "system", "display_name": "system", "role": "system"},
        )

    pool = _row(conn, "SELECT id FROM case_pools LIMIT 1")
    if pool:
        pool_id = pool["id"]
    else:
        pool_id = new_id()
        _insert(
            conn,
            "case_pools",
            {
                "id": pool_id,
                "name": "demo-pool",
                "purpose": "stage3",
                "owner_user_id": "system",
                "created_at": now,
            },
        )

    resolved_case_id = case_id or new_id()
    _insert(
        conn,
        "cases",
        {
            "id": resolved_case_id,
            "case_pool_id": pool_id,
            "name": "demo-case",
            "case_number": f"DEMO-{resolved_case_id[:8]}",
            "created_by": "system",
            "created_at": now,
        },
    )
    return resolved_case_id


def get_document(conn, document_id: str) -> Optional[dict[str, Any]]:
    return _row(conn, "SELECT * FROM documents WHERE id = ?", (document_id,))


def get_version(conn, version_id: str) -> Optional[dict[str, Any]]:
    return _row(conn, "SELECT * FROM document_versions WHERE id = ?", (version_id,))


def list_versions(conn, document_id: str) -> list[dict[str, Any]]:
    return _rows(
        conn,
        "SELECT * FROM document_versions WHERE document_id = ? ORDER BY version_no ASC",
        (document_id,),
    )


def list_pages(conn, version_id: str) -> list[dict[str, Any]]:
    return _rows(
        conn,
        "SELECT * FROM document_pages WHERE document_version_id = ? ORDER BY page_no ASC",
        (version_id,),
    )


def list_chunks(conn, version_id: str, active_only: bool = True) -> list[dict[str, Any]]:
    condition = " AND is_active = 1" if active_only else ""
    return _rows(
        conn,
        f"SELECT * FROM document_chunks WHERE document_version_id = ?{condition} ORDER BY ordinal ASC",
        (version_id,),
    )


def replace_chunks(conn, version_id: str, chunks: list[dict[str, Any]]) -> None:
    conn.execute("DELETE FROM document_chunks WHERE document_version_id = ?", (version_id,))
    for chunk in chunks:
        _insert(conn, "document_chunks", chunk)


def deactivate_version_chunks(conn, version_id: str) -> None:
    conn.execute(
        "UPDATE document_chunks SET is_active = 0, stale = 1 WHERE document_version_id = ?",
        (version_id,),
    )


def upsert_page(conn, fields: dict[str, Any]) -> None:
    existing = _row(
        conn,
        "SELECT id FROM document_pages WHERE document_version_id = ? AND page_no = ?",
        (fields["document_version_id"], fields["page_no"]),
    )
    if existing:
        _update(conn, "document_pages", existing["id"], {k: v for k, v in fields.items() if k != "id"})
    else:
        fields.setdefault("id", new_id())
        _insert(conn, "document_pages", fields)


def set_current_version(conn, document_id: str, version_id: str) -> None:
    conn.execute(
        "UPDATE document_versions SET is_current = 0 WHERE document_id = ?", (document_id,)
    )
    conn.execute("UPDATE document_versions SET is_current = 1 WHERE id = ?", (version_id,))
    _update(conn, "documents", document_id, {"current_version_id": version_id})


def next_version_no(conn, document_id: str) -> int:
    row = _row(
        conn,
        "SELECT COALESCE(MAX(version_no), 0) AS current FROM document_versions WHERE document_id = ?",
        (document_id,),
    )
    return int(row["current"]) + 1


def add_audit(conn, **fields) -> None:
    fields.setdefault("id", new_id())
    fields.setdefault("created_at", utc_now())
    fields.setdefault("detail_json", "{}")
    _insert(conn, "material_audit_events", fields)


# ---------- 授权（仅契约，完整 RBAC 未实现） ----------


def deny_all_auth(user_id: str | None, case_id: str | None, action: str) -> tuple[bool, str]:
    """默认策略：未接入真实授权前一律拒绝材料操作。"""
    return False, "authorization_not_configured"


def allow_all_auth(user_id: str | None, case_id: str | None, action: str) -> tuple[bool, str]:
    """本地与测试用桩，不是真实 RBAC。"""
    return True, "allow_all_stub"


def _default_auth():
    mode = os.getenv("MATERIAL_AUTH_MODE", "deny_all").lower()
    return allow_all_auth if mode == "allow_all" else deny_all_auth


# ---------- 敏感信息检测与脱敏 ----------

REDACTION_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    (
        "id_card",
        re.compile(
            r"(?<!\d)([1-9]\d{5}(?:19|20)\d{2}(?:0[1-9]|1[0-2])"
            r"(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx])(?!\d)"
        ),
    ),
    ("bank_card", re.compile(r"(?<!\d)([1-9]\d{15,18})(?!\d)")),
    ("phone", re.compile(r"(?<!\d)(1[3-9]\d{9})(?!\d)")),
    (
        "account",
        re.compile(r"(?i)(?:账号|帐户|账户|user(?:name)?|login)[:：\s]*([A-Za-z0-9_.-]{4,32})"),
    ),
    (
        "address",
        re.compile(
            r"([\u4e00-\u9fff]{2,10}(?:省|市|自治区|特别行政区))?"
            r"[\u4e00-\u9fff]{1,10}(?:市|州|盟)?"
            r"[\u4e00-\u9fff]{1,12}(?:区|县|旗)"
            r"[\u4e00-\u9fff0-9\-号弄幢栋单元室楼]{0,30}"
        ),
    ),
    ("name", re.compile(r"(?:姓名|被告人|嫌疑人|当事人)[:：\s]*([\u4e00-\u9fff·]{2,4})")),
]

_UNREDACTED_PATTERNS = [
    re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)"),
    re.compile(
        r"(?<!\d)[1-9]\d{5}(?:19|20)\d{2}(?:0[1-9]|1[0-2])"
        r"(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx](?!\d)"
    ),
]


@dataclass
class RedactionHit:
    sens_type: str
    start: int
    end: int
    original: str
    placeholder: str


def redact_text(text: str) -> tuple[str, list[RedactionHit]]:
    """返回脱敏文本与命中项；原值映射由 save_redaction_map 单独保管。"""
    hits: list[RedactionHit] = []
    for sensitive_type, pattern in REDACTION_PATTERNS:
        for match in pattern.finditer(text):
            if match.lastindex:
                start, end, original = match.start(1), match.end(1), match.group(1)
            else:
                start, end, original = match.start(), match.end(), match.group(0)
            if sensitive_type == "bank_card" and len(original) == 18:
                continue
            hits.append(RedactionHit(sensitive_type, start, end, original, placeholder=""))

    hits.sort(key=lambda hit: (hit.start, -(hit.end - hit.start)))
    accepted: list[RedactionHit] = []
    counters: dict[str, int] = {}
    occupied_until = -1
    for hit in hits:
        if hit.start < occupied_until:
            continue
        counters[hit.sens_type] = counters.get(hit.sens_type, 0) + 1
        hit.placeholder = f"[{hit.sens_type.upper()}_{counters[hit.sens_type]}]"
        accepted.append(hit)
        occupied_until = hit.end

    redacted = text
    for hit in sorted(accepted, key=lambda item: item.start, reverse=True):
        redacted = redacted[: hit.start] + hit.placeholder + redacted[hit.end :]
    return redacted, accepted


def save_redaction_map(version_id: str, items: list[dict[str, Any]], root=None) -> str:
    """原值与替换值的对照单独落盘，不进入分析载荷、日志或模型请求。"""
    directory = Path(root or REDACTION_STORAGE_DIR) / version_id
    directory.mkdir(parents=True, exist_ok=True)
    map_id = new_id()
    payload = {"map_id": map_id, "document_version_id": version_id, "items": items}
    (directory / f"{map_id}.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return map_id


def load_redaction_map(version_id: str, map_id: str, root=None) -> Optional[dict[str, Any]]:
    path = Path(root or REDACTION_STORAGE_DIR) / version_id / f"{map_id}.json"
    if not path.exists():
        return None
    return json.loads(path.read_text(encoding="utf-8"))


# ---------- 解析与页级 OCR ----------


@dataclass
class PageResult:
    page_no: int
    source: str
    text: str
    text_density: float = 0.0
    avg_confidence: float | None = None
    min_confidence: float | None = None
    bbox: list[Any] = field(default_factory=list)
    lines: list[dict[str, Any]] = field(default_factory=list)
    quality_flags: list[str] = field(default_factory=list)
    status: str = "PARSED"
    error_code: str | None = None
    error_message: str | None = None
    image_bytes: bytes | None = None


@dataclass
class OCRLine:
    text: str
    confidence: float
    bbox: list[float]


class OCREngine(Protocol):
    def recognize(self, image_bytes: bytes) -> list[OCRLine]:
        ...


class FallbackOCREngine:
    """未安装 PaddleOCR 时使用的确定性后备引擎，便于本地与测试运行。"""

    def recognize(self, image_bytes: bytes) -> list[OCRLine]:
        marker = b"__OCR_TEXT__:"
        if marker in image_bytes:
            text = image_bytes.split(marker, 1)[1].decode("utf-8", errors="ignore")
            return [OCRLine(text, 0.95 if text.strip() else 0.2, [0, 0, 100, 20])]

        printable = "".join(chr(value) if 32 <= value < 127 else " " for value in image_bytes)
        tokens = [token for token in printable.split() if len(token) >= 4]
        if not tokens:
            return [OCRLine("", 0.1, [0, 0, 1, 1])]
        return [OCRLine(" ".join(tokens[:50]), 0.55, [0, 0, 100, 20])]


class PaddleOCREngine:
    def __init__(self):
        from paddleocr import PaddleOCR

        self._ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)

    def recognize(self, image_bytes: bytes) -> list[OCRLine]:
        import numpy as np
        from PIL import Image

        image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        lines: list[OCRLine] = []
        for block in self._ocr.ocr(np.array(image), cls=True) or []:
            for box, (text, confidence) in block or []:
                flat = [float(value) for point in box for value in point]
                xs, ys = flat[0::2], flat[1::2]
                lines.append(
                    OCRLine(str(text), float(confidence), [min(xs), min(ys), max(xs), max(ys)])
                )
        return lines


_ocr_engine: OCREngine | None = None


def set_ocr_engine(engine: OCREngine | None) -> None:
    """注入 OCR 引擎，用于替换实现或测试。"""
    global _ocr_engine
    _ocr_engine = engine


def _get_ocr_engine() -> OCREngine:
    global _ocr_engine
    if _ocr_engine is None:
        try:
            _ocr_engine = PaddleOCREngine()
        except Exception:
            _ocr_engine = FallbackOCREngine()
    return _ocr_engine


def _recognize_image(image_bytes: bytes) -> dict[str, Any]:
    lines = _get_ocr_engine().recognize(image_bytes)
    texts = [line.text for line in lines if line.text]
    confidences = [line.confidence for line in lines] or [0.0]
    average = sum(confidences) / len(confidences)
    minimum = min(confidences)
    flags = []
    if minimum < OCR_LOW_CONFIDENCE_THRESHOLD or average < OCR_LOW_CONFIDENCE_THRESHOLD:
        flags.append("LOW_OCR_CONFIDENCE")
    if not "".join(texts).strip():
        flags.append("EMPTY_OCR")
    return {
        "text": "\n".join(texts),
        "lines": [
            {"text": line.text, "confidence": line.confidence, "bbox": line.bbox}
            for line in lines
        ],
        "avg_confidence": average,
        "min_confidence": minimum,
        "quality_flags": flags,
        "bbox": [line.bbox for line in lines],
    }


def _extract_pdf(path: Path) -> list[PageResult]:
    try:
        import fitz
    except ImportError as exc:
        raise MaterialError(ERROR_CODES["PARSE_FAILED"], "PyMuPDF(fitz) not installed") from exc

    try:
        document = fitz.open(path)
    except Exception as exc:
        message = str(exc).lower()
        code = (
            ERROR_CODES["ENCRYPTED_FILE"]
            if "password" in message or "encrypt" in message
            else ERROR_CODES["CORRUPT_FILE"]
        )
        raise MaterialError(code, f"Failed to open PDF: {exc}") from exc

    needs_password = getattr(document, "needs_pass", False) or getattr(
        document, "is_encrypted", False
    )
    if needs_password and not document.authenticate(""):
        document.close()
        raise MaterialError(ERROR_CODES["ENCRYPTED_FILE"], "PDF requires a password")

    pages: list[PageResult] = []
    try:
        for index, page in enumerate(document):
            try:
                text = page.get_text("text") or ""
                bbox = [
                    [float(block[0]), float(block[1]), float(block[2]), float(block[3])]
                    for block in (page.get_text("blocks") or [])
                    if len(block) >= 4
                ]
                area = float(page.rect.width * page.rect.height) if page.rect else 1.0
                density = len(text.strip()) / max(area / 10000.0, 1.0)
                needs_ocr = density < OCR_TEXT_DENSITY_THRESHOLD or not text.strip()
                pages.append(
                    PageResult(
                        page_no=index + 1,
                        source="pdf_text",
                        text=text,
                        text_density=density,
                        bbox=bbox,
                        status="NEEDS_OCR" if needs_ocr else "PARSED",
                        image_bytes=(
                            page.get_pixmap(matrix=fitz.Matrix(2, 2)).tobytes("png")
                            if needs_ocr
                            else None
                        ),
                        quality_flags=["LOW_TEXT_DENSITY"] if needs_ocr else [],
                    )
                )
            except Exception as exc:
                pages.append(
                    PageResult(
                        page_no=index + 1,
                        source="pdf_text",
                        text="",
                        status="FAILED",
                        error_code=ERROR_CODES["PARSE_FAILED"],
                        error_message=str(exc),
                        quality_flags=["PAGE_PARSE_FAILED"],
                    )
                )
    finally:
        document.close()
    return pages


def _extract_docx(path: Path) -> list[PageResult]:
    try:
        from docx import Document

        document = Document(str(path))
    except ImportError as exc:
        raise MaterialError(ERROR_CODES["PARSE_FAILED"], "python-docx not installed") from exc
    except Exception as exc:
        raise MaterialError(ERROR_CODES["CORRUPT_FILE"], f"Cannot open DOCX: {exc}") from exc

    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return [PageResult(1, "docx", "\n".join(parts), text_density=1.0)]


def _extract_txt(path: Path) -> list[PageResult]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            text = path.read_text(encoding="gbk")
        except Exception as exc:
            raise MaterialError(
                ERROR_CODES["CORRUPT_FILE"], f"Cannot decode text file: {exc}"
            ) from exc
    except Exception as exc:
        raise MaterialError(ERROR_CODES["CORRUPT_FILE"], f"Cannot read text file: {exc}") from exc
    return [PageResult(1, "txt", text, text_density=1.0)]


def _extract_image(path: Path) -> list[PageResult]:
    return [
        PageResult(
            1,
            "image",
            "",
            status="NEEDS_OCR",
            image_bytes=path.read_bytes(),
            quality_flags=["IMAGE_REQUIRES_OCR"],
        )
    ]


def _run_page_ocr(page: PageResult, retries: int) -> PageResult:
    if not page.image_bytes:
        page.status = "OCR_FAILED"
        page.error_code = ERROR_CODES["OCR_FAILED"]
        page.error_message = "missing page image for OCR"
        page.quality_flags = sorted(set(page.quality_flags + ["OCR_FAILED"]))
        return page

    last_error: Exception | None = None
    for _ in range(retries + 1):
        try:
            result = _recognize_image(page.image_bytes)
            page.text = result["text"]
            page.lines = result["lines"]
            page.bbox = result["bbox"]
            page.avg_confidence = result["avg_confidence"]
            page.min_confidence = result["min_confidence"]
            page.quality_flags = sorted(set(page.quality_flags + result["quality_flags"]))
            page.source = "ocr"
            if "EMPTY_OCR" in result["quality_flags"]:
                page.status = "OCR_FAILED"
                page.error_code = ERROR_CODES["OCR_FAILED"]
                page.error_message = "OCR returned empty text"
            elif "LOW_OCR_CONFIDENCE" in result["quality_flags"]:
                page.status = "NEEDS_OCR_REVIEW"
            else:
                page.status = "PARSED"
            return page
        except Exception as exc:
            last_error = exc

    page.status = "OCR_FAILED"
    page.error_code = ERROR_CODES["OCR_FAILED"]
    page.error_message = str(last_error) if last_error else "OCR failed"
    page.quality_flags = sorted(set(page.quality_flags + ["OCR_FAILED"]))
    return page


def parse_file_to_pages(path: Path | str, max_retries: int | None = None) -> list[dict[str, Any]]:
    """按文件类型取文，只对需要识别的页面执行 OCR，成功页不重跑。"""
    file_path = Path(path)
    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".txt": _extract_txt,
        ".png": _extract_image,
        ".jpg": _extract_image,
        ".jpeg": _extract_image,
    }
    extractor = extractors.get(file_path.suffix.lower())
    if extractor is None:
        raise MaterialError(
            ERROR_CODES["UNSUPPORTED_TYPE"], f"Unsupported extension: {file_path.suffix.lower()}"
        )

    retries = OCR_MAX_PAGE_RETRIES if max_retries is None else max_retries
    parsed = []
    for page in extractor(file_path):
        if page.status == "NEEDS_OCR":
            page = _run_page_ocr(page, retries)
        page.image_bytes = None
        parsed.append(
            {
                "page_no": page.page_no,
                "source": page.source,
                "text": page.text,
                "text_density": page.text_density,
                "avg_confidence": page.avg_confidence,
                "min_confidence": page.min_confidence,
                "bbox": page.bbox,
                "lines": page.lines,
                "quality_flags": page.quality_flags,
                "status": page.status,
                "error_code": page.error_code,
                "error_message": page.error_message,
            }
        )
    return parsed


def summarize_page_quality(pages: list[dict[str, Any]]) -> dict[str, Any]:
    """汇总页级识别质量，供状态展示与人工修正定位。"""
    confidences = [p["avg_confidence"] for p in pages if p.get("avg_confidence") is not None]
    minimums = [p["min_confidence"] for p in pages if p.get("min_confidence") is not None]
    low_pages: list[int] = []
    warnings: list[str] = []
    statuses: dict[str, str] = {}
    for page in pages:
        page_no = str(page["page_no"])
        status = page.get("status") or "PARSED"
        flags = page.get("quality_flags") or []
        statuses[page_no] = status
        if status in {"OCR_FAILED", "FAILED"}:
            warnings.append(f"page {page_no}: {status}")
        if "LOW_OCR_CONFIDENCE" in flags:
            low_pages.append(int(page["page_no"]))
            warnings.append(f"page {page_no}: low OCR confidence")
        if "EMPTY_OCR" in flags:
            warnings.append(f"page {page_no}: empty OCR")
    return {
        "avg_confidence": sum(confidences) / len(confidences) if confidences else None,
        "min_confidence": min(minimums) if minimums else None,
        "low_confidence_pages": sorted(set(low_pages)),
        "warnings": warnings,
        "page_statuses": statuses,
    }


def derive_version_status(pages: list[dict[str, Any]]) -> str:
    statuses = [page.get("status") for page in pages]
    if not statuses:
        return "FAILED"
    if any(status == "OCR_FAILED" for status in statuses):
        return "OCR_FAILED"
    if any(status == "NEEDS_OCR_REVIEW" for status in statuses):
        return "NEEDS_OCR_REVIEW"
    if any(status == "FAILED" for status in statuses):
        return "FAILED"
    return "PARSED" if all(status == "PARSED" for status in statuses) else "NEEDS_OCR_REVIEW"


# ---------- 分块与定位 ----------


def _text_sha256(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def build_chunks(
    document_version_id: str,
    pages: list[dict[str, Any]],
    *,
    chunk_size: int | None = None,
    overlap: int | None = None,
    parser_version: str | None = None,
) -> list[dict[str, Any]]:
    """生成稳定重叠文本块，ID 由版本、块序与文本哈希派生。"""
    size = CHUNK_SIZE if chunk_size is None else chunk_size
    overlap_size = CHUNK_OVERLAP if overlap is None else overlap
    version = parser_version or PARSER_VERSION
    if size <= 0:
        raise ValueError("chunk_size must be positive")
    if overlap_size >= size:
        overlap_size = max(0, size // 4)

    stream: list[tuple[str, int, list]] = []
    for page in pages:
        page_no = int(page["page_no"])
        text = page.get("text") or ""
        bbox = page.get("bbox") or []
        stream.extend((character, page_no, bbox) for character in text)
        if text and not text.endswith("\n"):
            stream.append(("\n", page_no, bbox))

    full_text = "".join(character for character, _, _ in stream)
    if not full_text.strip():
        digest = _text_sha256("")
        return [
            {
                "id": f"{document_version_id}:0:{digest[:16]}",
                "document_version_id": document_version_id,
                "ordinal": 0,
                "page_start": int(pages[0]["page_no"]) if pages else 1,
                "page_end": int(pages[-1]["page_no"]) if pages else 1,
                "char_start": 0,
                "char_end": 0,
                "bbox_json": "[]",
                "text_raw": "",
                "text_redacted": "",
                "text_sha256": digest,
                "parser_version": version,
                "quality_flags_json": "[]",
                "is_active": 1,
                "stale": 0,
            }
        ]

    chunks: list[dict[str, Any]] = []
    start = 0
    ordinal = 0
    while start < len(full_text):
        end = min(start + size, len(full_text))
        if end < len(full_text):
            window = full_text[start:end]
            break_at = max(window.rfind("\n"), window.rfind(" "), window.rfind("\u3000"))
            if break_at > size * 0.5:
                end = start + break_at + 1

        text = full_text[start:end]
        page_start = stream[start][1]
        page_end = stream[end - 1][1]
        flags: list[str] = []
        for page in pages:
            if page_start <= int(page["page_no"]) <= page_end:
                flags.extend(page.get("quality_flags") or [])
        digest = _text_sha256(text)
        chunks.append(
            {
                "id": f"{document_version_id}:{ordinal}:{digest[:16]}",
                "document_version_id": document_version_id,
                "ordinal": ordinal,
                "page_start": page_start,
                "page_end": page_end,
                "char_start": start,
                "char_end": end,
                "bbox_json": json.dumps(stream[start][2], ensure_ascii=False),
                "text_raw": text,
                "text_redacted": text,
                "text_sha256": digest,
                "parser_version": version,
                "quality_flags_json": json.dumps(sorted(set(flags)), ensure_ascii=False),
                "is_active": 1,
                "stale": 0,
            }
        )
        ordinal += 1
        if end >= len(full_text):
            break
        start = max(end - overlap_size, start + 1)
    return chunks


# ---------- 材料处理服务 ----------


class MaterialService:
    """案件卷宗上传后的处理入口：版本、解析、质量、修正、删除与安全读取。"""

    def __init__(self, db_path=None, auth_check=None, redaction_dir=None):
        self.db_path = db_path
        self.auth_check = auth_check or _default_auth()
        self.redaction_dir = Path(redaction_dir or REDACTION_STORAGE_DIR)
        init_db(self.db_path)

    def _authorize(self, user_id: str | None, case_id: str | None, action: str) -> None:
        allowed, reason = self.auth_check(user_id, case_id, action)
        if not allowed:
            with db_session(self.db_path) as conn:
                add_audit(
                    conn,
                    event_type="auth_check",
                    actor_user_id=user_id,
                    case_id=case_id,
                    decision="deny",
                    reason=reason or "denied",
                    detail_json=json.dumps({"action": action}, ensure_ascii=False),
                )
            raise MaterialError(ERROR_CODES["AUTH_DENIED"], reason or "authorization denied")

    # ----- 上传与版本 -----

    def upload_many(
        self,
        items: list[dict[str, Any]],
        *,
        user_id: str | None = None,
        parse: bool = True,
        keep_duplicate: bool = False,
    ) -> list[dict[str, Any]]:
        """一次导入多份材料，可分别归属不同案件。"""
        return [
            self.upload_one(
                case_id=item["case_id"],
                filename=item["filename"],
                content=item.get("content"),
                path=item.get("path"),
                user_id=user_id,
                parse=parse,
                keep_duplicate=keep_duplicate,
                replace_document_id=item.get("replace_document_id"),
            )
            for item in items
        ]

    def upload_one(
        self,
        *,
        case_id: str,
        filename: str,
        content: bytes | None = None,
        path: str | Path | None = None,
        user_id: str | None = None,
        parse: bool = True,
        keep_duplicate: bool = False,
        replace_document_id: str | None = None,
    ) -> dict[str, Any]:
        self._authorize(user_id, case_id, "material.upload")
        safe_name = Path(filename).name
        extension = Path(safe_name).suffix.lower()
        if extension not in ALLOWED_MATERIAL_EXTENSIONS:
            raise MaterialError(ERROR_CODES["UNSUPPORTED_TYPE"], f"unsupported type: {extension}")

        if content is None:
            if path is None:
                raise MaterialError(ERROR_CODES["CORRUPT_FILE"], "empty upload")
            content = Path(path).read_bytes()

        size = len(content)
        if size > MAX_UPLOAD_BYTES:
            raise MaterialError(
                ERROR_CODES["FILE_TOO_LARGE"],
                f"file exceeds limit {MAX_UPLOAD_BYTES} bytes",
                {"size": size},
            )
        if size == 0:
            raise MaterialError(ERROR_CODES["CORRUPT_FILE"], "empty file")

        digest = hashlib.sha256(content).hexdigest()
        content_type = mimetypes.guess_type(safe_name)[0] or "application/octet-stream"

        with db_session(self.db_path) as conn:
            if not get_case(conn, case_id):
                if user_id == "system":
                    ensure_demo_case(conn, case_id)
                else:
                    raise MaterialError(ERROR_CODES["NOT_FOUND"], f"case not found: {case_id}")

            duplicate = _row(
                conn,
                """
                SELECT d.* FROM documents d
                JOIN document_versions v ON v.document_id = d.id
                WHERE d.case_id = ? AND v.sha256 = ? AND d.deleted_at IS NULL
                ORDER BY v.created_at DESC LIMIT 1
                """,
                (case_id, digest),
            )
            if duplicate and not replace_document_id and not keep_duplicate:
                _update(conn, "documents", duplicate["id"], {"status": "DUPLICATE_PENDING"})
                add_audit(
                    conn,
                    event_type="duplicate_pending",
                    actor_user_id=user_id,
                    case_id=case_id,
                    document_id=duplicate["id"],
                    decision="pending",
                    reason="same sha256 in case",
                    detail_json=json.dumps({"sha256": digest}, ensure_ascii=False),
                )
                return {
                    "status": "DUPLICATE_PENDING",
                    "error_code": ERROR_CODES["DUPLICATE_PENDING"],
                    "document": get_document(conn, duplicate["id"]),
                    "message": "duplicate hash in case; decide keep or cancel",
                }

            now = utc_now()
            if replace_document_id:
                document = get_document(conn, replace_document_id)
                if not document or document["case_id"] != case_id:
                    raise MaterialError(ERROR_CODES["NOT_FOUND"], "document to replace not found")
                document_id = document["id"]
                parent_version_id = document.get("current_version_id")
                source_type = "REPLACE"
                if parent_version_id:
                    deactivate_version_chunks(conn, parent_version_id)
                    _update(conn, "document_versions", parent_version_id, {"is_current": 0})
            else:
                document_id = new_id()
                parent_version_id = None
                source_type = "UPLOAD"
                _insert(
                    conn,
                    "documents",
                    {
                        "id": document_id,
                        "case_id": case_id,
                        "filename": safe_name,
                        "stored_name": safe_name,
                        "content_type": content_type,
                        "size": size,
                        "sha256": digest,
                        "uploaded_by": user_id or "anonymous",
                        "created_at": now,
                        "status": "UPLOADED",
                        "current_version_id": None,
                        "deleted_at": None,
                        "quality_summary_json": "{}",
                    },
                )

            version_id = new_id()
            version_no = next_version_no(conn, document_id)
            storage_path = (
                Path(MATERIAL_STORAGE_DIR)
                / case_id
                / document_id
                / f"v{version_no}_{digest[:12]}{extension}"
            )
            storage_path.parent.mkdir(parents=True, exist_ok=True)
            storage_path.write_bytes(content)

            _insert(
                conn,
                "document_versions",
                {
                    "id": version_id,
                    "document_id": document_id,
                    "version_no": version_no,
                    "parent_version_id": parent_version_id,
                    "source_type": source_type,
                    "sha256": digest,
                    "storage_path": str(storage_path),
                    "content_type": content_type,
                    "size": size,
                    "parser_version": PARSER_VERSION,
                    "status": "UPLOADED",
                    "is_current": 1,
                    "is_active": 1,
                    "quality_summary_json": "{}",
                    "error_code": None,
                    "error_message": None,
                    "created_by": user_id or "anonymous",
                    "created_at": now,
                },
            )
            set_current_version(conn, document_id, version_id)
            _update(
                conn,
                "documents",
                document_id,
                {
                    "filename": safe_name,
                    "stored_name": storage_path.name,
                    "content_type": content_type,
                    "size": size,
                    "sha256": digest,
                    "status": "UPLOADED",
                },
            )
            add_audit(
                conn,
                event_type="upload",
                actor_user_id=user_id,
                case_id=case_id,
                document_id=document_id,
                document_version_id=version_id,
                decision="allow",
                reason=source_type.lower(),
            )

        if parse:
            return self.parse_version(version_id, user_id=user_id)
        with db_session(self.db_path) as conn:
            return {
                "document": get_document(conn, document_id),
                "version": get_version(conn, version_id),
                "status": "UPLOADED",
            }

    # ----- 解析 -----

    def parse_version(self, version_id: str, user_id: str | None = None) -> dict[str, Any]:
        with db_session(self.db_path) as conn:
            version = get_version(conn, version_id)
            if not version:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "version not found")
            document = get_document(conn, version["document_id"])
            case_id = document["case_id"] if document else None
            storage_path = Path(version["storage_path"])

        self._authorize(user_id, case_id, "material.parse")

        job_id = new_id()
        with db_session(self.db_path) as conn:
            now = utc_now()
            _insert(
                conn,
                "parse_jobs",
                {
                    "id": job_id,
                    "document_version_id": version_id,
                    "status": "RUNNING",
                    "attempt": 1,
                    "max_attempts": OCR_MAX_PAGE_RETRIES,
                    "error_code": None,
                    "error_message": None,
                    "started_at": now,
                    "finished_at": None,
                    "created_at": now,
                },
            )
            _update(conn, "document_versions", version_id, {"status": "PARSING"})
            if document:
                _update(conn, "documents", document["id"], {"status": "PARSING"})

        try:
            pages = parse_file_to_pages(storage_path)
        except Exception as exc:
            code = exc.code if isinstance(exc, MaterialError) else ERROR_CODES["PARSE_FAILED"]
            message = exc.message if isinstance(exc, MaterialError) else str(exc)
            with db_session(self.db_path) as conn:
                _update(
                    conn,
                    "parse_jobs",
                    job_id,
                    {
                        "status": "FAILED",
                        "error_code": code,
                        "error_message": message,
                        "finished_at": utc_now(),
                    },
                )
                _update(
                    conn,
                    "document_versions",
                    version_id,
                    {"status": "FAILED", "error_code": code, "error_message": message},
                )
                if document:
                    _update(conn, "documents", document["id"], {"status": "FAILED"})
            if isinstance(exc, MaterialError):
                raise
            raise MaterialError(code, message) from exc

        return self._persist_pages(
            version_id=version_id,
            document_id=version["document_id"],
            case_id=case_id,
            pages=pages,
            user_id=user_id,
            job_id=job_id,
            event_type="parse",
        )

    def _persist_pages(
        self,
        *,
        version_id: str,
        document_id: str,
        case_id: str | None,
        pages: list[dict[str, Any]],
        user_id: str | None,
        job_id: str | None,
        event_type: str,
        audit_reason: str | None = None,
    ) -> dict[str, Any]:
        """落库页面、分块与脱敏结果，并同步材料状态。"""
        quality = summarize_page_quality(pages)
        status = derive_version_status(pages)
        chunks = build_chunks(version_id, pages)

        map_items: list[dict[str, Any]] = []
        redaction_rows: list[dict[str, Any]] = []
        for chunk in chunks:
            redacted, hits = redact_text(chunk["text_raw"])
            chunk["text_redacted"] = redacted
            for hit in hits:
                map_items.append(
                    {
                        "chunk_id": chunk["id"],
                        "sens_type": hit.sens_type,
                        "start": hit.start,
                        "end": hit.end,
                        "original": hit.original,
                        "placeholder": hit.placeholder,
                    }
                )
                redaction_rows.append(
                    {
                        "id": new_id(),
                        "document_version_id": version_id,
                        "chunk_id": chunk["id"],
                        "sens_type": hit.sens_type,
                        "start_offset": hit.start,
                        "end_offset": hit.end,
                        "placeholder": hit.placeholder,
                        "map_ref": "",
                        "created_at": utc_now(),
                    }
                )

        map_id = save_redaction_map(version_id, map_items, root=self.redaction_dir)
        for row in redaction_rows:
            row["map_ref"] = map_id

        with db_session(self.db_path) as conn:
            for page in pages:
                upsert_page(
                    conn,
                    {
                        "id": new_id(),
                        "document_version_id": version_id,
                        "page_no": page["page_no"],
                        "source": page["source"],
                        "text": page["text"],
                        "text_density": page.get("text_density") or 0,
                        "avg_confidence": page.get("avg_confidence"),
                        "min_confidence": page.get("min_confidence"),
                        "bbox_json": json.dumps(page.get("bbox") or [], ensure_ascii=False),
                        "lines_json": json.dumps(page.get("lines") or [], ensure_ascii=False),
                        "quality_flags_json": json.dumps(
                            page.get("quality_flags") or [], ensure_ascii=False
                        ),
                        "status": page.get("status") or "PARSED",
                        "retry_count": 0,
                        "error_code": page.get("error_code"),
                        "error_message": page.get("error_message"),
                    },
                )
            replace_chunks(conn, version_id, chunks)
            conn.execute(
                "DELETE FROM redaction_items WHERE document_version_id = ?", (version_id,)
            )
            for row in redaction_rows:
                _insert(conn, "redaction_items", row)

            quality_json = json.dumps(quality, ensure_ascii=False)
            _update(
                conn,
                "document_versions",
                version_id,
                {"status": status, "quality_summary_json": quality_json},
            )
            _update(
                conn,
                "documents",
                document_id,
                {"status": status, "quality_summary_json": quality_json},
            )
            if job_id:
                _update(conn, "parse_jobs", job_id, {"status": "DONE", "finished_at": utc_now()})
            add_audit(
                conn,
                event_type=event_type,
                actor_user_id=user_id,
                case_id=case_id,
                document_id=document_id,
                document_version_id=version_id,
                decision="allow",
                reason=audit_reason or status,
            )
            return {
                "document": get_document(conn, document_id),
                "version": get_version(conn, version_id),
                "pages": list_pages(conn, version_id),
                "chunks": [
                    {key: value for key, value in chunk.items() if key != "text_raw"}
                    for chunk in list_chunks(conn, version_id)
                ],
                "quality": quality,
                "status": status,
            }

    # ----- 查询 -----

    def list_materials(self, case_id: str, user_id: str | None = None) -> list[dict[str, Any]]:
        self._authorize(user_id, case_id, "material.list")
        with db_session(self.db_path) as conn:
            documents = _rows(
                conn,
                "SELECT * FROM documents WHERE case_id = ? AND deleted_at IS NULL ORDER BY created_at DESC",
                (case_id,),
            )
            materials = []
            for document in documents:
                versions = list_versions(conn, document["id"])
                materials.append(
                    {
                        **document,
                        "quality_summary": json.loads(document.get("quality_summary_json") or "{}"),
                        "current_version": next(
                            (v for v in versions if v.get("is_current")), None
                        ),
                        "version_count": len(versions),
                    }
                )
            return materials

    def get_status(self, document_id: str, user_id: str | None = None) -> dict[str, Any]:
        with db_session(self.db_path) as conn:
            document = get_document(conn, document_id)
            if not document:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "document not found")
            case_id = document["case_id"]

        self._authorize(user_id, case_id, "material.read")

        with db_session(self.db_path) as conn:
            document = get_document(conn, document_id)
            current = (
                get_version(conn, document["current_version_id"])
                if document.get("current_version_id")
                else None
            )
            pages = list_pages(conn, current["id"]) if current else []
            low_quality = []
            for page in pages:
                flags = json.loads(page.get("quality_flags_json") or "[]")
                if page["status"] in {"NEEDS_OCR_REVIEW", "OCR_FAILED", "FAILED"} or (
                    "LOW_OCR_CONFIDENCE" in flags
                ):
                    low_quality.append(
                        {
                            "page_no": page["page_no"],
                            "status": page["status"],
                            "quality_flags": flags,
                            "min_confidence": page.get("min_confidence"),
                            "avg_confidence": page.get("avg_confidence"),
                        }
                    )
            return {
                "document": document,
                "versions": list_versions(conn, document_id),
                "current_version": current,
                "quality_summary": json.loads(document.get("quality_summary_json") or "{}"),
                "low_quality_pages": low_quality,
            }

    # ----- 外发门控 -----

    def read_redacted_chunk(
        self,
        document_version_id: str,
        chunk_id: str | None = None,
        user_id: str | None = None,
    ) -> dict[str, Any]:
        """经门控读取脱敏片段：版本有效、已脱敏、最小必要片段。"""
        with db_session(self.db_path) as conn:
            version = get_version(conn, document_version_id)
            if not version:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "version not found")
            document = get_document(conn, version["document_id"])
            case_id = document["case_id"] if document else None

        self._authorize(user_id, case_id, "material.read_redacted")

        with db_session(self.db_path) as conn:
            version = get_version(conn, document_version_id)
            document_id = version["document_id"]

            def deny(reason: str, message: str, code: str, detail: dict | None = None):
                add_audit(
                    conn,
                    event_type="egress_check",
                    actor_user_id=user_id,
                    case_id=case_id,
                    document_id=document_id,
                    document_version_id=document_version_id,
                    decision="deny",
                    reason=reason,
                    detail_json=json.dumps(detail or {}, ensure_ascii=False),
                )
                return MaterialError(code, message)

            if not version.get("is_active"):
                raise deny(
                    "version_inactive",
                    "inactive version cannot egress",
                    ERROR_CODES["EGRESS_DENIED"],
                )

            if chunk_id:
                chunk = _row(conn, "SELECT * FROM document_chunks WHERE id = ?", (chunk_id,))
            else:
                available = list_chunks(conn, document_version_id, active_only=True)
                chunk = available[0] if available else None
            if not chunk:
                raise deny("chunk_not_found", "chunk not found", ERROR_CODES["NOT_FOUND"])

            if not chunk.get("is_active") or chunk.get("stale"):
                raise deny(
                    "chunk_stale_or_inactive",
                    "stale/inactive chunk cannot egress",
                    ERROR_CODES["EGRESS_DENIED"],
                    {"chunk_id": chunk["id"]},
                )

            redactions = _rows(
                conn,
                "SELECT * FROM redaction_items WHERE document_version_id = ?",
                (document_version_id,),
            )
            unchanged = chunk.get("text_raw") and chunk["text_redacted"] == chunk["text_raw"]
            if unchanged and redactions:
                raise deny(
                    "unredacted_with_items",
                    "unredacted text cannot egress",
                    ERROR_CODES["EGRESS_DENIED"],
                    {"chunk_id": chunk["id"]},
                )
            text = chunk.get("text_redacted") or ""
            if any(pattern.search(text) for pattern in _UNREDACTED_PATTERNS):
                raise deny(
                    "looks_unredacted",
                    "unredacted sensitive text cannot egress",
                    ERROR_CODES["EGRESS_DENIED"],
                    {"chunk_id": chunk["id"]},
                )

            add_audit(
                conn,
                event_type="egress_check",
                actor_user_id=user_id,
                case_id=case_id,
                document_id=document_id,
                document_version_id=document_version_id,
                decision="allow",
                reason="ok",
                detail_json=json.dumps({"chunk_id": chunk["id"]}, ensure_ascii=False),
            )
            return {
                "chunk_id": chunk["id"],
                "document_version_id": document_version_id,
                "ordinal": chunk["ordinal"],
                "page_start": chunk["page_start"],
                "page_end": chunk["page_end"],
                "text": chunk["text_redacted"],
                "quality_flags": json.loads(chunk.get("quality_flags_json") or "[]"),
                "redacted": True,
            }

    # ----- 人工修正 -----

    def apply_correction(
        self,
        *,
        document_id: str,
        source_version_id: str,
        page_no: int,
        corrected_text: str,
        user_id: str | None = None,
    ) -> dict[str, Any]:
        """人工修正生成新版本，旧版本原文与 chunk 保留但标记失效。"""
        with db_session(self.db_path) as conn:
            document = get_document(conn, document_id)
            if not document:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "document not found")
            case_id = document["case_id"]

        self._authorize(user_id, case_id, "material.correct")

        with db_session(self.db_path) as conn:
            source = get_version(conn, source_version_id)
            if not source or source["document_id"] != document_id:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "source version not found")
            pages = list_pages(conn, source_version_id)
            if not pages:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "source pages not found")

            corrected_pages: list[dict[str, Any]] = []
            found = False
            for page in pages:
                item = {
                    "page_no": page["page_no"],
                    "source": page["source"],
                    "text": page["text"],
                    "text_density": page["text_density"],
                    "avg_confidence": page["avg_confidence"],
                    "min_confidence": page["min_confidence"],
                    "bbox": json.loads(page.get("bbox_json") or "[]"),
                    "lines": json.loads(page.get("lines_json") or "[]"),
                    "quality_flags": json.loads(page.get("quality_flags_json") or "[]"),
                    "status": "PARSED",
                    "error_code": None,
                    "error_message": None,
                }
                if int(page["page_no"]) == int(page_no):
                    found = True
                    item.update(
                        {
                            "text": corrected_text,
                            "source": "human_correction",
                            "avg_confidence": 1.0,
                            "min_confidence": 1.0,
                            "quality_flags": ["HUMAN_CORRECTED"],
                        }
                    )
                corrected_pages.append(item)
            if not found:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], f"page {page_no} not found")

            version_id = new_id()
            deactivate_version_chunks(conn, source_version_id)
            _update(conn, "document_versions", source_version_id, {"is_current": 0})
            _insert(
                conn,
                "document_versions",
                {
                    "id": version_id,
                    "document_id": document_id,
                    "version_no": next_version_no(conn, document_id),
                    "parent_version_id": source_version_id,
                    "source_type": "CORRECTION",
                    "sha256": source["sha256"],
                    "storage_path": source["storage_path"],
                    "content_type": source["content_type"],
                    "size": source["size"],
                    "parser_version": PARSER_VERSION,
                    "status": "PARSING",
                    "is_current": 1,
                    "is_active": 1,
                    "quality_summary_json": "{}",
                    "error_code": None,
                    "error_message": None,
                    "created_by": user_id or "anonymous",
                    "created_at": utc_now(),
                },
            )
            set_current_version(conn, document_id, version_id)

        result = self._persist_pages(
            version_id=version_id,
            document_id=document_id,
            case_id=case_id,
            pages=corrected_pages,
            user_id=user_id,
            job_id=None,
            event_type="correction",
            audit_reason=f"correct_page_{page_no}",
        )
        result["parent_version_id"] = source_version_id
        return result

    # ----- 删除与重复裁决 -----

    def preview_delete_impact(self, document_id: str, user_id: str | None = None) -> dict[str, Any]:
        with db_session(self.db_path) as conn:
            document = get_document(conn, document_id)
            if not document:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "document not found")
            case_id = document["case_id"]

        self._authorize(user_id, case_id, "material.delete")

        with db_session(self.db_path) as conn:
            impact = []
            for version in list_versions(conn, document_id):
                active = _row(
                    conn,
                    "SELECT COUNT(*) AS total FROM document_chunks WHERE document_version_id = ? AND is_active = 1",
                    (version["id"],),
                )["total"]
                impact.append(
                    {
                        "kind": "version",
                        "id": version["id"],
                        "detail": f"v{version['version_no']} status={version['status']} active_chunks={active}",
                    }
                )
                for chunk in list_chunks(conn, version["id"], active_only=False):
                    impact.append(
                        {
                            "kind": "chunk",
                            "id": chunk["id"],
                            "detail": f"ordinal={chunk['ordinal']} stale={chunk['stale']}",
                        }
                    )
            return {"document_id": document_id, "impact": impact}

    def logical_delete(self, document_id: str, user_id: str | None = None) -> dict[str, Any]:
        impact = self.preview_delete_impact(document_id, user_id=user_id)
        with db_session(self.db_path) as conn:
            document = get_document(conn, document_id)
            _update(
                conn,
                "documents",
                document_id,
                {"deleted_at": utc_now(), "status": "DELETED"},
            )
            for version in list_versions(conn, document_id):
                deactivate_version_chunks(conn, version["id"])
                _update(
                    conn, "document_versions", version["id"], {"is_active": 0, "is_current": 0}
                )
            add_audit(
                conn,
                event_type="logical_delete",
                actor_user_id=user_id,
                case_id=document["case_id"],
                document_id=document_id,
                decision="allow",
                reason="logical_delete",
                detail_json=json.dumps(impact, ensure_ascii=False),
            )
            return {"status": "DELETED", **impact}

    def resolve_duplicate(
        self, document_id: str, *, action: str, user_id: str | None = None
    ) -> dict[str, Any]:
        """裁决重复导入：keep 保留并解析，cancel 逻辑删除。"""
        with db_session(self.db_path) as conn:
            document = get_document(conn, document_id)
            if not document:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "document not found")
            case_id = document["case_id"]
            current_version_id = document.get("current_version_id")

        self._authorize(user_id, case_id, "material.upload")

        if action == "cancel":
            return self.logical_delete(document_id, user_id=user_id)
        if action == "keep":
            if not current_version_id:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "no current version")
            with db_session(self.db_path) as conn:
                _update(conn, "documents", document_id, {"status": "UPLOADED"})
            return self.parse_version(current_version_id, user_id=user_id)
        raise MaterialError(ERROR_CODES["PARSE_FAILED"], f"unknown action: {action}")


_default_service: MaterialService | None = None


def get_material_service() -> MaterialService:
    global _default_service
    if _default_service is None:
        _default_service = MaterialService()
    return _default_service


def set_material_service(service: MaterialService | None) -> None:
    global _default_service
    _default_service = service


# ---------- Agent 工具（只返回脱敏内容） ----------


def _dump(data: Any) -> str:
    return json.dumps(data, ensure_ascii=False, indent=2)


def list_case_materials(case_id: str, user_id: Optional[str] = None) -> str:
    """列出某案件下材料及处理状态、质量摘要。"""
    try:
        materials = get_material_service().list_materials(case_id, user_id=user_id or "system")
        return _dump(
            {
                "materials": [
                    {
                        "document_id": material["id"],
                        "filename": material["filename"],
                        "status": material["status"],
                        "quality_summary": material.get("quality_summary"),
                        "current_version_id": material.get("current_version_id"),
                        "version_count": material.get("version_count"),
                    }
                    for material in materials
                ]
            }
        )
    except MaterialError as exc:
        return _dump(exc.to_dict())


def get_material_status(document_id: str, user_id: Optional[str] = None) -> str:
    """查询单份材料处理状态、版本链与低质量页。"""
    try:
        return _dump(get_material_service().get_status(document_id, user_id=user_id or "system"))
    except MaterialError as exc:
        return _dump(exc.to_dict())


def locate_low_quality_pages(document_id: str, user_id: Optional[str] = None) -> str:
    """定位识别质量不佳、需要人工修正的页面。"""
    try:
        status = get_material_service().get_status(document_id, user_id=user_id or "system")
        return _dump(
            {
                "document_id": document_id,
                "low_quality_pages": status.get("low_quality_pages", []),
                "quality_summary": status.get("quality_summary"),
            }
        )
    except MaterialError as exc:
        return _dump(exc.to_dict())


def read_material_chunk(
    document_version_id: str,
    chunk_id: Optional[str] = None,
    user_id: Optional[str] = None,
) -> str:
    """经外发门控读取已脱敏的材料片段。"""
    try:
        return _dump(
            get_material_service().read_redacted_chunk(
                document_version_id, chunk_id=chunk_id, user_id=user_id or "system"
            )
        )
    except MaterialError as exc:
        return _dump(exc.to_dict())


def submit_ocr_correction(
    document_id: str,
    source_version_id: str,
    page_no: int,
    corrected_text: str,
    user_id: Optional[str] = None,
) -> str:
    """提交 OCR 人工修正，生成新版本且不覆盖历史。"""
    try:
        return _dump(
            get_material_service().apply_correction(
                document_id=document_id,
                source_version_id=source_version_id,
                page_no=int(page_no),
                corrected_text=corrected_text,
                user_id=user_id or "system",
            )
        )
    except MaterialError as exc:
        return _dump(exc.to_dict())
