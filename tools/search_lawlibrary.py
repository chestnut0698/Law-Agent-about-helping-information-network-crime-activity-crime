import os
import re
import json
import glob
from typing import Optional, List, Dict


class DocxLawLibrary:
    """本地法律法规 docx 库

    私有成员以 _ 开头。
    功能：
    1. 扫描指定目录下的所有 .docx 法律文件
    2. 解析为 {法规名: {条号: 条文内容}}
    3. 支持按法规名 + 条号精确提取
    4. 支持关键词检索
    5. 为智能体提供结构化引用
    """

    def __init__(self, laws_dir: str = "./laws"):
        self._laws_dir = laws_dir
        self._cache = {}  # {文件名: {title, clauses: {条号: 内容}}}
        self._index = {}  # {法规名关键词: 文件名}

        # 条文匹配正则：匹配"第一条"、"第五百七十七条"等
        self._clause_pattern = re.compile(
            r'^第[一二三四五六七八九十百千零0-9]+条'
        )
        # 章节匹配正则
        self._chapter_pattern = re.compile(
            r'^第[一二三四五六七八九十]+章|^第一编|^第二编|^第三编|^第四编|^第五编|^第六编|^第七编'
        )

        self._build_index()

    # ---------- 私有方法 ----------
    def _human_delay(self, min_s: float = 0.1, max_s: float = 0.3):
        import time
        import random
        time.sleep(random.uniform(min_s, max_s))

    def _extract_text_from_docx(self, file_path: str) -> str:
        """提取 docx 全文文本（含表格）"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请先安装 python-docx: pip install python-docx")

        doc = Document(file_path)
        texts = []

        # 1. 提取所有段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)

        # 2. 提取所有表格的单元格文本（关键！法律 docx 常在表格里）
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        texts.append(cell_text)

        return "\n".join(texts)

    def _parse_clauses(self, full_text: str) -> Dict[str, str]:
        """将全文解析为 {条号: 条文内容}"""
        lines = full_text.split("\n")
        clauses = {}
        current_clause = None
        buffer = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 匹配条号
            match = self._clause_pattern.match(line)
            if match:
                # 保存上一条
                if current_clause and buffer:
                    clauses[current_clause] = "\n".join(buffer).strip()

                # 开始新一条
                current_clause = match.group(0)
                # 条文可能在同一行： "第一条 为了保护..."
                rest = line[len(current_clause):].strip()
                # 去掉可能的【标题】部分
                rest = re.sub(r'^【.*?】', '', rest).strip()
                buffer = [rest] if rest else []
            else:
                # 非条号行，追加到当前条文
                if current_clause:
                    # 跳过章节标题行
                    if not self._chapter_pattern.match(line):
                        buffer.append(line)

        # 保存最后一条
        if current_clause and buffer:
            clauses[current_clause] = "\n".join(buffer).strip()

        return clauses

    def _build_index(self):
        """扫描 laws_dir，建立索引"""
        if not os.path.exists(self._laws_dir):
            os.makedirs(self._laws_dir, exist_ok=True)
            return

        for ext in ("*.docx", "*.DOCX"):
            for file_path in glob.glob(os.path.join(self._laws_dir, ext)):
                fname = os.path.basename(file_path)
                try:
                    full_text = self._extract_text_from_docx(file_path)
                    clauses = self._parse_clauses(full_text)

                    # 从文件名推测法规标题（去掉 .docx 和后缀）
                    title = re.sub(r'\.docx$', '', fname, flags=re.I)
                    title = re.sub(r'【.*?】', '', title).strip()  # 去掉【2018年修订】等

                    self._cache[fname] = {
                        "title": title,
                        "file_path": file_path,
                        "total_clauses": len(clauses),
                        "clauses": clauses,
                        "text_length": len(full_text),
                    }

                    # 建立标题关键词 → 文件名 的反向索引
                    # 用标题的每个 2-gram 作为关键词
                    for i in range(len(title) - 1):
                        bigram = title[i:i + 2]
                        self._index.setdefault(bigram, []).append(fname)

                    # 完整标题也作为关键词
                    self._index.setdefault(title, []).append(fname)

                except Exception as e:
                    print(f"⚠️ 解析 {fname} 失败: {e}")

    def _find_law_file(self, title_keyword: str) -> Optional[str]:
        """根据关键词找到匹配的法规文件名"""
        # 1. 精确匹配
        if title_keyword in self._index:
            return self._index[title_keyword][0]

        # 2. 子串匹配（反向查 cache）
        for fname, info in self._cache.items():
            if title_keyword in info["title"]:
                return fname

        # 3. bigram 匹配
        best_match = None
        best_score = 0
        for i in range(len(title_keyword) - 1):
            bigram = title_keyword[i:i + 2]
            if bigram in self._index:
                for fname in self._index[bigram]:
                    # 计算标题相似度
                    score = 0
                    cache_title = self._cache[fname]["title"]
                    if title_keyword in cache_title:
                        score = len(title_keyword)
                    elif any(bigram in cache_title for bigram in
                             [title_keyword[i:i + 2] for i in range(len(title_keyword) - 1)]):
                        score = sum(1 for bg in
                                    [title_keyword[i:i + 2] for i in range(len(title_keyword) - 1)]
                                    if bg in cache_title)
                    if score > best_score:
                        best_score = score
                        best_match = fname

        return best_match

    # ---------- 公开方法 ----------
    def list_laws(self) -> List[str]:
        """列出库中所有法规标题"""
        return [info["title"] for info in self._cache.values()]

    def get_clause(self, law_title: str, clause_num: str = "") -> Dict:
        """获取指定法规的条文

        Args:
            law_title: 法规标题关键词（如"民法典"）
            clause_num: 条号（如"第一条"），留空返回目录
        """
        fname = self._find_law_file(law_title)
        if not fname:
            return {
                "found": False,
                "available_laws": self.list_laws()[:10],  # 返回前10个供参考
                "message": f"未找到包含'{law_title}'的法规，可用法规见 available_laws"
            }

        info = self._cache[fname]
        clauses = info["clauses"]

        if not clause_num:
            # 返回目录
            return {
                "found": True,
                "title": info["title"],
                "total_clauses": info["total_clauses"],
                "toc": list(clauses.keys()),
            }

        # 精确匹配条号
        if clause_num in clauses:
            return {
                "found": True,
                "title": info["title"],
                "clause": clause_num,
                "text": clauses[clause_num],
            }

        # 模糊匹配（如"五百七十七"匹配"第五百七十七条"）
        for k, v in clauses.items():
            # 提取条号中的数字部分进行比较
            if clause_num in k:
                return {
                    "found": True,
                    "title": info["title"],
                    "clause": k,
                    "text": v,
                }

        # 用阿拉伯数字匹配
        try:
            arabic_num = int(clause_num)
            chinese_num = self._arabic_to_chinese(arabic_num)
            target = f"第{chinese_num}条"
            if target in clauses:
                return {
                    "found": True,
                    "title": info["title"],
                    "clause": target,
                    "text": clauses[target],
                }
        except ValueError:
            pass

        return {
            "found": True,
            "title": info["title"],
            "clause_requested": clause_num,
            "text": f"未找到第'{clause_num}'条，该法规共有 {info['total_clauses']} 条。",
        }

    def search_by_keyword(self, keyword: str, law_title: str = "") -> Dict:
        """关键词检索条文

        Args:
            keyword: 检索关键词（如"合同"）
            law_title: 限定法规范围（可选）
        """
        results = []

        # 确定搜索范围
        if law_title:
            fname = self._find_law_file(law_title)
            files_to_search = [fname] if fname else []
        else:
            files_to_search = list(self._cache.keys())

        for fname in files_to_search:
            info = self._cache[fname]
            for clause_num, content in info["clauses"].items():
                if keyword in content or keyword in clause_num:
                    results.append({
                        "law_title": info["title"],
                        "clause": clause_num,
                        "text": content,
                    })

        return {
            "keyword": keyword,
            "total_matches": len(results),
            "results": results[:20],  # 限制返回数量
        }

    def get_law_context(self, law_title: str, max_chars: int = 12000) -> str:
        """获取法规全文（截断），用于拼接到 LLM context"""
        fname = self._find_law_file(law_title)
        if not fname:
            return ""

        info = self._cache[fname]
        parts = []
        total = 0
        for clause_num, content in info["clauses"].items():
            block = f"【{clause_num}】{content}\n"
            if total + len(block) > max_chars:
                break
            parts.append(block)
            total += len(block)

        return f"《{info['title']}》（共{info['total_clauses']}条，节选前{total}字符）:\n" + "\n".join(parts)

    def reload(self):
        """重新扫描目录（新增 docx 后调用）"""
        self._cache.clear()
        self._index.clear()
        self._build_index()

    def _arabic_to_chinese(self, num: int) -> str:
        """阿拉伯数字转中文数字（用于条号匹配）"""
        digits = "零一二三四五六七八九"
        units = ["", "十", "百", "千"]
        if num == 0:
            return "零"
        if num < 10:
            return digits[num]
        if num < 20:
            return "十" + (digits[num % 10] if num % 10 != 0 else "")
        result = ""
        s = str(num)
        for i, ch in enumerate(s):
            digit = int(ch)
            unit = units[len(s) - i - 1]
            if digit == 0:
                if not result.endswith("零"):
                    result += "零"
            else:
                result += digits[digit] + unit
        return result.rstrip("零")

    # ---------- LLM 工具入口 ----------
    def list_laws_for_llm(self) -> str:
        """LLM 工具：列出所有可用法规"""
        return json.dumps({"laws": self.list_laws()}, ensure_ascii=False)

    def get_clause_for_llm(self, law_title: str, clause: str = "") -> str:
        """LLM 工具：获取条文"""
        return json.dumps(self.get_clause(law_title, clause), ensure_ascii=False)

    def search_for_llm(self, keyword: str, law_title: str = "") -> str:
        """LLM 工具：关键词检索"""
        return json.dumps(self.search_by_keyword(keyword, law_title), ensure_ascii=False)


def search_lawlibrary(action: str, law_title: str = "", clause: str = "", keyword: str = "") -> str:
    """
    统一的法律法规查询工具函数。
    参数：
        action: "list" | "get_clause" | "search"
        law_title: 法规标题关键词（get_clause/search 时需要）
        clause: 条号（get_clause 时可选，留空返回目录）
        keyword: 检索关键词（search 时需要）
        laws_dir: docx 文件所在目录（默认 ./laws）
    返回：JSON 字符串
    """
    laws_dir = "../knowledge/legal_library"
    lib = DocxLawLibrary(laws_dir=laws_dir)

    if action == "list":
        return json.dumps({"laws": lib.list_laws()}, ensure_ascii=False)

    elif action == "get_clause":
        result = lib.get_clause(law_title, clause)
        return json.dumps(result, ensure_ascii=False)

    elif action == "search":
        result = lib.search_by_keyword(keyword, law_title)
        return json.dumps(result, ensure_ascii=False)

    else:
        return json.dumps({"error": f"未知 action: {action}"}, ensure_ascii=False)