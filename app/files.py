"""案件电子卷宗的文件上传与后续处理。

覆盖上传落盘与哈希、版本链、解析与页级 OCR、质量状态、分块、脱敏、
外发门控、人工修正，以及提供给 Agent 调用的材料工具。
"""

from __future__ import annotations

import hashlib
import io
import json
import mimetypes
import re
import sqlite3
import uuid
from contextlib import contextmanager
import contextvars
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterator, Optional, Protocol, List, Tuple
from prikit import PDFAnonymizer
from presidio_analyzer import AnalyzerEngine
from presidio_analyzer.nlp_engine import NlpEngineProvider
from presidio_analyzer import RecognizerResult

from app.config import DATABASE_PATH, MATERIAL_STORAGE_DIR, REDACTION_STORAGE_DIR

_db_connection_ctx = contextvars.ContextVar('_db_connection_ctx', default=None)

# ---------- 常量 ----------
PRIKIT_ENTITY_MAP = {
    "PERSON": "name",
    "PHONE_NUMBER": "phone",
    "ID": "id_card",
    "CREDIT_CARD": "bank_card",
    "LOCATION": "address",
    "EMAIL_ADDRESS": "email",
    "IP_ADDRESS": "ip",
    "URL": "url",
    "DATE_TIME": "datetime",
    "NRP": "nrp",
    "CRYPTO": "crypto",
    "IBAN_CODE": "iban",
}

CONFIDENCE_THRESHOLD = {
    "name": 0.7,
    "address": 0.6,
    "phone": 0.7,
    "id_card": 0.7,
    "bank_card": 0.7,
    "email": 0.9,
    "ip": 0.9,
}

ALLOWED_MATERIAL_EXTENSIONS = {".pdf", ".png", ".jpg", ".jpeg", ".docx", ".txt"}
CHUNK_OVERLAP = 120
CHUNK_SIZE = 1000
MATERIAL_AUTH_MODE = "allow_all"
MAX_UPLOAD_BYTES = 50 * 1024 * 1024
OCR_TEXT_DENSITY_THRESHOLD = 0.08
OCR_LOW_CONFIDENCE_THRESHOLD = 0.75
OCR_MAX_PAGE_RETRIES = 2
PARSER_VERSION = "stage3-v1"

MATERIAL_STATUSES = {
    "UPLOADED",
    "PARSING",
    "PARSED",
    "NEEDS_OCR_REVIEW",
    "OCR_FAILED",
    "DUPLICATE_PENDING",
    "FAILED",
    "DELETED",
}

ERROR_CODES = {
    "UNSUPPORTED_TYPE": "MATERIAL_UNSUPPORTED_TYPE",
    "FILE_TOO_LARGE": "MATERIAL_FILE_TOO_LARGE",
    "CORRUPT_FILE": "MATERIAL_CORRUPT_FILE",
    "ENCRYPTED_FILE": "MATERIAL_ENCRYPTED_FILE",
    "OCR_FAILED": "MATERIAL_OCR_FAILED",
    "PARSE_FAILED": "MATERIAL_PARSE_FAILED",
    "AUTH_DENIED": "MATERIAL_AUTH_DENIED",
    "EGRESS_DENIED": "MATERIAL_EGRESS_DENIED",
    "NOT_FOUND": "MATERIAL_NOT_FOUND",
    "DUPLICATE_PENDING": "MATERIAL_DUPLICATE_PENDING",
    "CITATION_STALE": "MATERIAL_CITATION_STALE",
}


class MaterialError(Exception):
    def __init__(self, code: str, message: str, details: dict | None = None):
        self.code = code
        self.message = message
        self.details = details or {}
        super().__init__(f"{code}: {message}")

    def to_dict(self) -> dict[str, Any]:
        return {"error_code": self.code, "message": self.message, "details": self.details}


# ---------- 数据库 ----------
SCHEMA_SQL = """
CREATE TABLE IF NOT EXISTS users (
    id VARCHAR(64) PRIMARY KEY,
    display_name VARCHAR(100) NOT NULL,
    role VARCHAR(32) NOT NULL
);

CREATE TABLE IF NOT EXISTS case_pools (
    id VARCHAR(36) PRIMARY KEY,
    name VARCHAR(120) NOT NULL,
    purpose TEXT NOT NULL,
    owner_user_id VARCHAR(64) NOT NULL,
    created_at DATETIME NOT NULL
);

CREATE TABLE IF NOT EXISTS case_pool_members (
    id VARCHAR(36) PRIMARY KEY,
    case_pool_id VARCHAR(36) NOT NULL,
    user_id VARCHAR(64) NOT NULL,
    member_role VARCHAR(32) NOT NULL,
    UNIQUE(case_pool_id, user_id)
);

CREATE TABLE IF NOT EXISTS cases (
    id VARCHAR(36) PRIMARY KEY,
    case_pool_id VARCHAR(36) NOT NULL,
    name VARCHAR(160) NOT NULL,
    case_number VARCHAR(100) NOT NULL,
    created_by VARCHAR(64) NOT NULL,
    created_at DATETIME NOT NULL,
    UNIQUE(case_pool_id, case_number)
);

CREATE TABLE IF NOT EXISTS documents (
    id VARCHAR(36) PRIMARY KEY,
    case_id VARCHAR(36) NOT NULL,
    filename VARCHAR(255) NOT NULL,
    stored_name VARCHAR(255) NOT NULL,
    content_type VARCHAR(120) NOT NULL,
    size BIGINT NOT NULL,
    sha256 VARCHAR(64) NOT NULL,
    uploaded_by VARCHAR(64) NOT NULL,
    created_at DATETIME NOT NULL,
    status VARCHAR(32) NOT NULL DEFAULT 'UPLOADED',
    current_version_id VARCHAR(36),
    deleted_at DATETIME,
    quality_summary_json TEXT NOT NULL DEFAULT '{}'
);

CREATE TABLE IF NOT EXISTS document_versions (
    id VARCHAR(36) PRIMARY KEY,
    document_id VARCHAR(36) NOT NULL,
    version_no INTEGER NOT NULL,
    parent_version_id VARCHAR(36),
    source_type VARCHAR(32) NOT NULL,
    sha256 VARCHAR(64) NOT NULL,
    storage_path TEXT NOT NULL,
    content_type VARCHAR(120) NOT NULL,
    size BIGINT NOT NULL,
    parser_version VARCHAR(64) NOT NULL,
    status VARCHAR(32) NOT NULL,
    is_current INTEGER NOT NULL DEFAULT 0,
    is_active INTEGER NOT NULL DEFAULT 1,
    quality_summary_json TEXT NOT NULL DEFAULT '{}',
    error_code VARCHAR(64),
    error_message TEXT,
    created_by VARCHAR(64) NOT NULL,
    created_at DATETIME NOT NULL,
    UNIQUE(document_id, version_no)
);

CREATE TABLE IF NOT EXISTS parse_jobs (
    id VARCHAR(36) PRIMARY KEY,
    document_version_id VARCHAR(36) NOT NULL,
    status VARCHAR(32) NOT NULL,
    attempt INTEGER NOT NULL DEFAULT 0,
    max_attempts INTEGER NOT NULL DEFAULT 2,
    error_code VARCHAR(64),
    error_message TEXT,
    started_at DATETIME,
    finished_at DATETIME,
    created_at DATETIME NOT NULL
);

CREATE TABLE IF NOT EXISTS document_pages (
    id VARCHAR(36) PRIMARY KEY,
    document_version_id VARCHAR(36) NOT NULL,
    page_no INTEGER NOT NULL,
    source VARCHAR(32) NOT NULL,
    text TEXT NOT NULL DEFAULT '',
    text_density REAL NOT NULL DEFAULT 0,
    avg_confidence REAL,
    min_confidence REAL,
    bbox_json TEXT NOT NULL DEFAULT '[]',
    lines_json TEXT NOT NULL DEFAULT '[]',
    quality_flags_json TEXT NOT NULL DEFAULT '[]',
    status VARCHAR(32) NOT NULL,
    retry_count INTEGER NOT NULL DEFAULT 0,
    error_code VARCHAR(64),
    error_message TEXT,
    UNIQUE(document_version_id, page_no)
);

CREATE TABLE IF NOT EXISTS document_chunks (
    id VARCHAR(64) PRIMARY KEY,
    document_version_id VARCHAR(36) NOT NULL,
    ordinal INTEGER NOT NULL,
    page_start INTEGER NOT NULL,
    page_end INTEGER NOT NULL,
    char_start INTEGER NOT NULL,
    char_end INTEGER NOT NULL,
    bbox_json TEXT NOT NULL DEFAULT '[]',
    text_raw TEXT NOT NULL,
    text_redacted TEXT NOT NULL,
    text_sha256 VARCHAR(64) NOT NULL,
    parser_version VARCHAR(64) NOT NULL,
    quality_flags_json TEXT NOT NULL DEFAULT '[]',
    is_active INTEGER NOT NULL DEFAULT 1,
    stale INTEGER NOT NULL DEFAULT 0,
    UNIQUE(document_version_id, ordinal)
);

CREATE TABLE IF NOT EXISTS redaction_items (
    id VARCHAR(36) PRIMARY KEY,
    document_version_id VARCHAR(36) NOT NULL,
    chunk_id VARCHAR(64),
    sens_type VARCHAR(32) NOT NULL,
    start_offset INTEGER NOT NULL,
    end_offset INTEGER NOT NULL,
    placeholder VARCHAR(64) NOT NULL,
    map_ref VARCHAR(128) NOT NULL,
    created_at DATETIME NOT NULL
);

CREATE TABLE IF NOT EXISTS material_audit_events (
    id VARCHAR(36) PRIMARY KEY,
    event_type VARCHAR(64) NOT NULL,
    actor_user_id VARCHAR(64),
    case_id VARCHAR(36),
    document_id VARCHAR(36),
    document_version_id VARCHAR(36),
    decision VARCHAR(32) NOT NULL,
    reason TEXT NOT NULL,
    detail_json TEXT NOT NULL DEFAULT '{}',
    created_at DATETIME NOT NULL
);

CREATE TABLE IF NOT EXISTS entity_global_map (
    fingerprint VARCHAR(64) PRIMARY KEY,
    anonymous_id VARCHAR(48) NOT NULL UNIQUE,
    sens_type VARCHAR(24) NOT NULL,
    task_id VARCHAR(36) NOT NULL DEFAULT '',
    first_seen_at DATETIME NOT NULL,
    last_seen_at DATETIME NOT NULL
);

CREATE INDEX IF NOT EXISTS idx_documents_case ON documents(case_id);
CREATE INDEX IF NOT EXISTS idx_versions_document ON document_versions(document_id);
CREATE INDEX IF NOT EXISTS idx_pages_version ON document_pages(document_version_id);
CREATE INDEX IF NOT EXISTS idx_chunks_version ON document_chunks(document_version_id);
CREATE INDEX IF NOT EXISTS idx_redaction_version ON redaction_items(document_version_id);
"""

_global_mapper_instance: GlobalEntityMapper | None = None
_global_analyzer_instance: AnalyzerEngine | None = None


def get_global_mapper(db_path=None, salt: str = "default-salt-change-me") -> GlobalEntityMapper:
    global _global_mapper_instance
    if _global_mapper_instance is None:
        _global_mapper_instance = GlobalEntityMapper(db_path=db_path, salt=salt)
    return _global_mapper_instance


def get_global_analyzer() -> AnalyzerEngine:
    global _global_analyzer_instance
    if _global_analyzer_instance is None:
        configuration = {
            "nlp_engine_name": "spacy",
            "models": [{"lang_code": "zh", "model_name": "zh_core_web_trf"}],
        }
        provider = NlpEngineProvider(nlp_configuration=configuration)
        nlp_engine = provider.create_engine()
        _global_analyzer_instance = AnalyzerEngine(nlp_engine=nlp_engine)
    return _global_analyzer_instance


def new_id() -> str:
    return str(uuid.uuid4())


def utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def get_connection(db_path=None):
    # 如果上下文中已有连接，直接返回（复用）
    conn = _db_connection_ctx.get()
    if conn is not None:
        return conn
    # 否则创建新连接
    conn = sqlite3.connect(str(db_path or DATABASE_PATH), check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    conn.execute("PRAGMA journal_mode=WAL")
    return conn


def init_db(db_path=None) -> None:
    conn = get_connection(db_path)
    try:
        conn.executescript(SCHEMA_SQL)
        existing = {row["name"] for row in conn.execute("PRAGMA table_info(documents)").fetchall()}
        for column, ddl in (
            ("status", "status VARCHAR(32) NOT NULL DEFAULT 'UPLOADED'"),
            ("current_version_id", "current_version_id VARCHAR(36)"),
            ("deleted_at", "deleted_at DATETIME"),
            ("quality_summary_json", "quality_summary_json TEXT NOT NULL DEFAULT '{}'"),
        ):
            if column not in existing:
                conn.execute(f"ALTER TABLE documents ADD COLUMN {ddl}")
        conn.commit()
    finally:
        conn.close()


@contextmanager
def db_session(db_path=None):
    # 检查是否已有上下文连接，有则复用，无则新建
    existing_conn = _db_connection_ctx.get()
    if existing_conn is not None:
        yield existing_conn
        return
    # 没有上下文连接，新建
    conn = get_connection(db_path)
    token = _db_connection_ctx.set(conn)
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        _db_connection_ctx.reset(token)
        conn.close()


def _row(conn, sql: str, params=()) -> Optional[dict[str, Any]]:
    row = conn.execute(sql, params).fetchone()
    return dict(row) if row else None


def _rows(conn, sql: str, params=()) -> list[dict[str, Any]]:
    return [dict(row) for row in conn.execute(sql, params).fetchall()]


def _insert(conn, table: str, fields: dict[str, Any]) -> None:
    columns = list(fields)
    placeholders = ",".join("?" for _ in columns)
    conn.execute(
        f"INSERT INTO {table}({','.join(columns)}) VALUES ({placeholders})",
        [fields[column] for column in columns],
    )


def _update(conn, table: str, row_id: str, fields: dict[str, Any]) -> None:
    if not fields:
        return
    assignments = ", ".join(f"{key} = ?" for key in fields)
    conn.execute(
        f"UPDATE {table} SET {assignments} WHERE id = ?",
        [*fields.values(), row_id],
    )


def get_case(conn, case_id: str) -> Optional[dict[str, Any]]:
    return _row(conn, "SELECT * FROM cases WHERE id = ?", (case_id,))


def ensure_demo_case(conn, case_id: str | None = None) -> str:
    if case_id and get_case(conn, case_id):
        return case_id
    if case_id is None:
        existing = _row(conn, "SELECT id FROM cases LIMIT 1")
        if existing:
            return existing["id"]

    now = utc_now()
    if not _row(conn, "SELECT id FROM users WHERE id = ?", ("system",)):
        _insert(
            conn,
            "users",
            {"id": "system", "display_name": "system", "role": "system"},
        )

    pool = _row(conn, "SELECT id FROM case_pools LIMIT 1")
    if pool:
        pool_id = pool["id"]
    else:
        pool_id = new_id()
        _insert(
            conn,
            "case_pools",
            {
                "id": pool_id,
                "name": "demo-pool",
                "purpose": "stage3",
                "owner_user_id": "system",
                "created_at": now,
            },
        )

    resolved_case_id = case_id or new_id()
    _insert(
        conn,
        "cases",
        {
            "id": resolved_case_id,
            "case_pool_id": pool_id,
            "name": "demo-case",
            "case_number": f"DEMO-{resolved_case_id[:8]}",
            "created_by": "system",
            "created_at": now,
        },
    )
    return resolved_case_id


def get_document(conn, document_id: str) -> Optional[dict[str, Any]]:
    return _row(conn, "SELECT * FROM documents WHERE id = ?", (document_id,))


def get_version(conn, version_id: str) -> Optional[dict[str, Any]]:
    return _row(conn, "SELECT * FROM document_versions WHERE id = ?", (version_id,))


def list_versions(conn, document_id: str) -> list[dict[str, Any]]:
    return _rows(
        conn,
        "SELECT * FROM document_versions WHERE document_id = ? ORDER BY version_no ASC",
        (document_id,),
    )


def list_pages(conn, version_id: str) -> list[dict[str, Any]]:
    return _rows(
        conn,
        "SELECT * FROM document_pages WHERE document_version_id = ? ORDER BY page_no ASC",
        (version_id,),
    )


def list_chunks(conn, version_id: str, active_only: bool = True) -> list[dict[str, Any]]:
    condition = " AND is_active = 1" if active_only else ""
    return _rows(
        conn,
        f"SELECT * FROM document_chunks WHERE document_version_id = ?{condition} ORDER BY ordinal ASC",
        (version_id,),
    )


def replace_chunks(conn, version_id: str, chunks: list[dict[str, Any]]) -> None:
    conn.execute("DELETE FROM document_chunks WHERE document_version_id = ?", (version_id,))
    for chunk in chunks:
        _insert(conn, "document_chunks", chunk)


def deactivate_version_chunks(conn, version_id: str) -> None:
    conn.execute(
        "UPDATE document_chunks SET is_active = 0, stale = 1 WHERE document_version_id = ?",
        (version_id,),
    )


def upsert_page(conn, fields: dict[str, Any]) -> None:
    existing = _row(
        conn,
        "SELECT id FROM document_pages WHERE document_version_id = ? AND page_no = ?",
        (fields["document_version_id"], fields["page_no"]),
    )
    if existing:
        _update(conn, "document_pages", existing["id"], {k: v for k, v in fields.items() if k != "id"})
    else:
        fields.setdefault("id", new_id())
        _insert(conn, "document_pages", fields)


def set_current_version(conn, document_id: str, version_id: str) -> None:
    conn.execute(
        "UPDATE document_versions SET is_current = 0 WHERE document_id = ?", (document_id,)
    )
    conn.execute("UPDATE document_versions SET is_current = 1 WHERE id = ?", (version_id,))
    _update(conn, "documents", document_id, {"current_version_id": version_id})


def next_version_no(conn, document_id: str) -> int:
    row = _row(
        conn,
        "SELECT COALESCE(MAX(version_no), 0) AS current FROM document_versions WHERE document_id = ?",
        (document_id,),
    )
    return int(row["current"]) + 1


def add_audit(conn, **fields) -> None:
    fields.setdefault("id", new_id())
    fields.setdefault("created_at", utc_now())
    fields.setdefault("detail_json", "{}")
    _insert(conn, "material_audit_events", fields)


# ---------- 授权 ----------
def deny_all_auth(user_id: str | None, case_id: str | None, action: str) -> tuple[bool, str]:
    return False, "未配置材料授权。本地演示请在 .env 设置 MATERIAL_AUTH_MODE=allow_all 后重启服务"


def allow_all_auth(user_id: str | None, case_id: str | None, action: str) -> tuple[bool, str]:
    return True, "allow_all_stub"


def _default_auth():
    mode = (MATERIAL_AUTH_MODE or "").strip().lower() or "allow_all"
    return allow_all_auth if mode == "allow_all" else deny_all_auth


# ---------- GlobalEntityMapper ----------
class GlobalEntityMapper:
    def __init__(self, db_path=None, salt: str = "default-salt-change-me"):
        self.db_path = db_path
        self.salt = salt

    def _fingerprint(self, original: str) -> str:
        return hashlib.sha256(f"{original}{self.salt}".encode()).hexdigest()

    def _new_anonymous_id(self, sens_type: str) -> str:
        short_uuid = uuid.uuid4().hex[:8]
        return f"{sens_type.upper()}_{short_uuid}"

    def _is_in_context(self, conn) -> bool:
        return _db_connection_ctx.get() is conn

    def delete_by_task_id(self, task_id: str) -> int:
        conn = get_connection(self.db_path)
        try:
            cursor = conn.execute(
                "DELETE FROM entity_global_map WHERE task_id = ?",
                (task_id,)
            )
            if not self._is_in_context(conn):
                conn.commit()
            return cursor.rowcount
        finally:
            if not self._is_in_context(conn):
                conn.close()

    def get_or_create(self, original: str, sens_type: str, task_id: str = "") -> str:
        fp = self._fingerprint(original)
        conn = get_connection(self.db_path)
        row = conn.execute(
            "SELECT anonymous_id FROM entity_global_map WHERE fingerprint = ?",
            (fp,)
        ).fetchone()
        if row:
            conn.execute(
                "UPDATE entity_global_map SET last_seen_at = ?, task_id = COALESCE(NULLIF(task_id,''), ?) WHERE fingerprint = ?",
                (utc_now(), task_id, fp)
            )
            if not self._is_in_context(conn):
                conn.commit()
            return row[0]

        anon_id = self._new_anonymous_id(sens_type)
        now = utc_now()
        conn.execute(
            "INSERT INTO entity_global_map (fingerprint, anonymous_id, sens_type, task_id, first_seen_at, last_seen_at) "
            "VALUES (?, ?, ?, ?, ?, ?)",
            (fp, anon_id, sens_type, task_id, now, now)
        )
        if not self._is_in_context(conn):
            conn.commit()
        return anon_id

    def get_fingerprint(self, original: str) -> str:
        return self._fingerprint(original)

    def get_fingerprint_by_anonymous_id(self, anonymous_id: str) -> str | None:
        conn = get_connection(self.db_path)
        try:
            row = conn.execute(
                "SELECT fingerprint FROM entity_global_map WHERE anonymous_id = ?",
                (anonymous_id,)
            ).fetchone()
            return row[0] if row else None
        finally:
            if not self._is_in_context(conn):
                conn.close()

    def list_mappings(
            self,
            task_id: str | None = None,
            sens_type: str | None = None,
            anonymous_id: str | None = None,
            limit: int = 100,
            offset: int = 0
    ) -> dict[str, Any]:
        conn = get_connection(self.db_path)
        try:
            query = """
                SELECT fingerprint, anonymous_id, sens_type, task_id, first_seen_at, last_seen_at
                FROM entity_global_map
                WHERE 1=1
            """
            params = []
            if task_id is not None and task_id != "":
                query += " AND task_id = ?"
                params.append(task_id)
            if sens_type:
                query += " AND sens_type = ?"
                params.append(sens_type)
            if anonymous_id:
                query += " AND anonymous_id = ?"
                params.append(anonymous_id)

            count_query = f"SELECT COUNT(*) as total FROM ({query})"
            total_row = conn.execute(count_query, params).fetchone()
            total = total_row[0] if total_row else 0

            query += " ORDER BY last_seen_at DESC LIMIT ? OFFSET ?"
            params.extend([limit, offset])
            rows = conn.execute(query, params).fetchall()
            items = [dict(row) for row in rows]

            for item in items:
                fingerprint = item["fingerprint"]
                redact_row = conn.execute(
                    """
                    SELECT document_version_id, chunk_id, start_offset, end_offset
                    FROM redaction_items
                    WHERE map_ref = ?
                    LIMIT 1
                    """,
                    (fingerprint,)
                ).fetchone()
                if redact_row:
                    chunk_row = conn.execute(
                        "SELECT text_raw FROM document_chunks WHERE id = ?",
                        (redact_row["chunk_id"],)
                    ).fetchone()
                    if chunk_row:
                        text_raw = chunk_row["text_raw"] or ""
                        start = redact_row["start_offset"]
                        end = redact_row["end_offset"]
                        if 0 <= start < end <= len(text_raw):
                            item["sample_raw"] = text_raw[start:end]
                        else:
                            item["sample_raw"] = None
                    else:
                        item["sample_raw"] = None
                else:
                    item["sample_raw"] = None

            return {
                "total": total,
                "limit": limit,
                "offset": offset,
                "items": items
            }
        finally:
            if not self._is_in_context(conn):
                conn.close()

    def update_mapping(
            self,
            fingerprint: str,
            new_anonymous_id: str | None = None,
            new_sens_type: str | None = None
    ) -> bool:
        if not new_anonymous_id and not new_sens_type:
            return False
        conn = get_connection(self.db_path)
        try:
            updates = []
            params = []
            if new_anonymous_id is not None:
                updates.append("anonymous_id = ?")
                params.append(new_anonymous_id)
            if new_sens_type is not None:
                updates.append("sens_type = ?")
                params.append(new_sens_type)
            params.append(fingerprint)
            sql = f"UPDATE entity_global_map SET {', '.join(updates)} WHERE fingerprint = ?"
            conn.execute(sql, params)
            if not self._is_in_context(conn):
                conn.commit()
            return True
        finally:
            if not self._is_in_context(conn):
                conn.close()

    def delete_mapping(self, fingerprint: str) -> tuple[bool, str]:
        conn = get_connection(self.db_path)
        try:
            # 检查引用（如果有引用则拒绝删除）
            ref = conn.execute(
                "SELECT COUNT(*) FROM redaction_items WHERE map_ref = ?",
                (fingerprint,)
            ).fetchone()[0]
            if ref > 0:
                return False, f"该映射被 {ref} 条脱敏记录引用，无法删除"
            # ★ 关键：执行真正的 DELETE
            conn.execute("DELETE FROM entity_global_map WHERE fingerprint = ?", (fingerprint,))
            if not self._is_in_context(conn):
                conn.commit()
            return True, "删除成功"
        finally:
            if not self._is_in_context(conn):
                conn.close()

    # --------------------------------------------------------------
    # 核心重构方法：batch_apply_and_redact
    # --------------------------------------------------------------
    def batch_apply_and_redact(self, payload: dict) -> dict[str, Any]:
        """
        批量应用变更（删除、更新、新增映射），并重脱敏所有受影响的 chunk。
        所有操作在一个事务中完成，基于原文和现有 redaction_items 重建脱敏文本，
        不调用 redact_text。
        """
        with db_session(self.db_path) as conn:
            deletions = payload.get("deletions", [])
            updates = payload.get("updates", {})
            additions = payload.get("additions", [])
            document_id = payload.get("document_id")

            # ---------- 1. 应用删除和更新（影响 redaction_items） ----------
            affected_chunks = {}

            # 收集删除影响的 chunk
            for fp in deletions:
                if document_id:
                    rows = conn.execute(
                        """
                        SELECT ri.id, ri.chunk_id, ri.start_offset, ri.end_offset, ri.placeholder, ri.map_ref
                        FROM redaction_items ri
                        JOIN document_chunks dc ON dc.id = ri.chunk_id
                        JOIN document_versions dv ON dv.id = dc.document_version_id
                        WHERE ri.map_ref = ? AND dv.document_id = ?
                        """,
                        (fp, document_id)
                    ).fetchall()
                else:
                    rows = conn.execute(
                        "SELECT id, chunk_id, start_offset, end_offset, placeholder, map_ref FROM redaction_items WHERE map_ref = ?",
                        (fp,)
                    ).fetchall()
                for row in rows:
                    affected_chunks.setdefault(row["chunk_id"], []).append(dict(row))

            # 收集更新影响的 chunk
            for fp in updates:
                if document_id:
                    rows = conn.execute(
                        """
                        SELECT ri.id, ri.chunk_id, ri.start_offset, ri.end_offset, ri.placeholder, ri.map_ref
                        FROM redaction_items ri
                        JOIN document_chunks dc ON dc.id = ri.chunk_id
                        JOIN document_versions dv ON dv.id = dc.document_version_id
                        WHERE ri.map_ref = ? AND dv.document_id = ?
                        """,
                        (fp, document_id)
                    ).fetchall()
                else:
                    rows = conn.execute(
                        "SELECT id, chunk_id, start_offset, end_offset, placeholder, map_ref FROM redaction_items WHERE map_ref = ?",
                        (fp,)
                    ).fetchall()
                for row in rows:
                    affected_chunks.setdefault(row["chunk_id"], []).append(dict(row))

            # 执行更新映射（entity_global_map 层面）
            for fp, update in updates.items():
                sets = []
                params = []
                if "sens_type" in update:
                    sets.append("sens_type = ?")
                    params.append(update["sens_type"])
                if "anonymous_id" in update:
                    sets.append("anonymous_id = ?")
                    params.append(update["anonymous_id"])
                if sets:
                    params.append(fp)
                    conn.execute(
                        f"UPDATE entity_global_map SET {', '.join(sets)} WHERE fingerprint = ?",
                        params
                    )

            # 对受影响的 chunk 进行重脱敏（删除和更新）
            redacted_count = 0
            for chunk_id, old_items in affected_chunks.items():
                chunk = conn.execute(
                    "SELECT text_raw, document_version_id FROM document_chunks WHERE id = ?",
                    (chunk_id,)
                ).fetchone()
                if not chunk:
                    continue

                text_raw = chunk["text_raw"]
                doc_version_id = chunk["document_version_id"]
                new_text = text_raw
                new_items = []

                # 按 start_offset 降序
                sorted_items = sorted(old_items, key=lambda x: x["start_offset"], reverse=True)

                for item in sorted_items:
                    start = item["start_offset"]
                    end = item["end_offset"]
                    map_ref = item["map_ref"]

                    if map_ref in deletions:
                        placeholder = text_raw[start:end]  # 恢复原文
                    else:
                        # 检查更新
                        update = updates.get(map_ref)
                        if update:
                            new_anon = conn.execute(
                                "SELECT anonymous_id, sens_type FROM entity_global_map WHERE fingerprint = ?",
                                (map_ref,)
                            ).fetchone()
                            if new_anon:
                                new_placeholder = new_anon["anonymous_id"]
                                sens_type = new_anon["sens_type"]
                                new_items.append({
                                    "id": new_id(),
                                    "document_version_id": doc_version_id,
                                    "chunk_id": chunk_id,
                                    "sens_type": sens_type,
                                    "start_offset": start,
                                    "end_offset": end,
                                    "placeholder": new_placeholder,
                                    "map_ref": map_ref,
                                    "created_at": utc_now()
                                })
                                placeholder = new_placeholder
                            else:
                                placeholder = text_raw[start:end]
                        else:
                            # 保持不变
                            cur_anon = conn.execute(
                                "SELECT anonymous_id, sens_type FROM entity_global_map WHERE fingerprint = ?",
                                (map_ref,)
                            ).fetchone()
                            if cur_anon:
                                new_items.append({
                                    "id": new_id(),
                                    "document_version_id": doc_version_id,
                                    "chunk_id": chunk_id,
                                    "sens_type": cur_anon["sens_type"],
                                    "start_offset": start,
                                    "end_offset": end,
                                    "placeholder": cur_anon["anonymous_id"],
                                    "map_ref": map_ref,
                                    "created_at": utc_now()
                                })
                                placeholder = cur_anon["anonymous_id"]
                            else:
                                placeholder = text_raw[start:end]

                    new_text = new_text[:start] + placeholder + new_text[end:]

                # 更新 chunk
                conn.execute(
                    "UPDATE document_chunks SET text_redacted = ? WHERE id = ?",
                    (new_text, chunk_id)
                )
                conn.execute("DELETE FROM redaction_items WHERE chunk_id = ?", (chunk_id,))
                for item_data in new_items:
                    conn.execute(
                        """
                        INSERT INTO redaction_items 
                        (id, document_version_id, chunk_id, sens_type, start_offset, end_offset, placeholder, map_ref, created_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            item_data["id"],
                            item_data["document_version_id"],
                            item_data["chunk_id"],
                            item_data["sens_type"],
                            item_data["start_offset"],
                            item_data["end_offset"],
                            item_data["placeholder"],
                            item_data["map_ref"],
                            item_data["created_at"]
                        )
                    )
                redacted_count += 1

            # ---------- 2. 处理新增映射（先插入映射，再应用到所有文档） ----------
            for add in additions:
                original = add.get("original")
                sens_type = add.get("sens_type")
                if not original or not sens_type:
                    continue
                fp = self._fingerprint(original)
                existing = conn.execute(
                    "SELECT fingerprint FROM entity_global_map WHERE fingerprint = ?",
                    (fp,)
                ).fetchone()
                if existing:
                    # 映射已存在，可能不需要重复插入，但我们可以继续应用
                    # 仍然需要获取 anonymous_id
                    anon_info = conn.execute(
                        "SELECT anonymous_id, sens_type FROM entity_global_map WHERE fingerprint = ?",
                        (fp,)
                    ).fetchone()
                    if not anon_info:
                        continue
                    placeholder = anon_info["anonymous_id"]
                    sens_type = anon_info["sens_type"]
                else:
                    anon_id = self._new_anonymous_id(sens_type)
                    now = utc_now()
                    conn.execute(
                        """
                        INSERT INTO entity_global_map 
                        (fingerprint, anonymous_id, sens_type, task_id, first_seen_at, last_seen_at)
                        VALUES (?, ?, ?, ?, ?, ?)
                        """,
                        (fp, anon_id, sens_type, '', now, now)
                    )
                    placeholder = anon_id

                # 如果指定了文档，只处理该文档；否则处理所有文档
                if document_id:
                    chunk_rows = conn.execute(
                        """
                        SELECT dc.id, dc.text_raw, dc.document_version_id
                        FROM document_chunks dc
                        JOIN document_versions dv ON dv.id = dc.document_version_id
                        WHERE dv.document_id = ? AND dv.is_current = 1 AND dv.is_active = 1 AND dc.is_active = 1
                        """,
                        (document_id,)
                    ).fetchall()
                else:
                    # 全局生效：处理所有文档的所有当前版本 chunk
                    chunk_rows = conn.execute(
                        """
                        SELECT dc.id, dc.text_raw, dc.document_version_id
                        FROM document_chunks dc
                        JOIN document_versions dv ON dv.id = dc.document_version_id
                        WHERE dv.is_current = 1 AND dv.is_active = 1 AND dc.is_active = 1
                        """
                    ).fetchall()

                for chunk_row in chunk_rows:
                    chunk_id = chunk_row["id"]
                    text_raw = chunk_row["text_raw"]
                    doc_version_id = chunk_row["document_version_id"]

                    # 获取现有 redaction_items 的占用区间
                    existing_items = conn.execute(
                        "SELECT start_offset, end_offset FROM redaction_items WHERE chunk_id = ?",
                        (chunk_id,)
                    ).fetchall()
                    occupied_intervals = [(row["start_offset"], row["end_offset"]) for row in existing_items]

                    # 查找所有匹配位置（非重叠）
                    import re
                    positions = []
                    start_pos = 0
                    while True:
                        idx = text_raw.find(original, start_pos)
                        if idx == -1:
                            break
                        positions.append((idx, idx + len(original)))
                        start_pos = idx + 1

                    # 筛选未被占用的位置
                    new_positions = []
                    for start, end in positions:
                        overlapped = False
                        for os, oe in occupied_intervals:
                            if not (end <= os or start >= oe):
                                overlapped = True
                                break
                        if not overlapped:
                            new_positions.append((start, end))

                    if not new_positions:
                        continue

                    # 获取当前 text_redacted
                    cur_redacted = conn.execute(
                        "SELECT text_redacted FROM document_chunks WHERE id = ?",
                        (chunk_id,)
                    ).fetchone()[0]

                    # 按 start 降序，从后往前替换
                    new_positions_sorted = sorted(new_positions, key=lambda x: x[0], reverse=True)
                    new_text = cur_redacted
                    for start, end in new_positions_sorted:
                        new_text = new_text[:start] + placeholder + new_text[end:]

                    # 更新 chunk
                    conn.execute(
                        "UPDATE document_chunks SET text_redacted = ? WHERE id = ?",
                        (new_text, chunk_id)
                    )

                    # 插入新的 redaction_items
                    for start, end in new_positions_sorted:
                        conn.execute(
                            """
                            INSERT INTO redaction_items 
                            (id, document_version_id, chunk_id, sens_type, start_offset, end_offset, placeholder, map_ref, created_at)
                            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
                            """,
                            (
                                new_id(),
                                doc_version_id,
                                chunk_id,
                                sens_type,
                                start,
                                end,
                                placeholder,
                                fp,
                                utc_now()
                            )
                        )
                    redacted_count += 1

            # ---------- 3. 删除映射（entity_global_map） ----------
            for fp in deletions:
                conn.execute("DELETE FROM entity_global_map WHERE fingerprint = ?", (fp,))

            return {
                "ok": True,
                "deleted": len(deletions),
                "updated": len(updates),
                "added": len(additions),
                "redacted_chunks": redacted_count,
                "message": f"变更已应用，重脱敏 {redacted_count} 个文本块"
            }

# ---------- 敏感信息检测与脱敏 ----------
REDACTION_PATTERNS: list[tuple[str, re.Pattern[str]]] = [
    ("id_card", re.compile(
        r"(?<!\d)([1-9]\d{5}(?:19|20)\d{2}(?:0[1-9]|1[0-2])"
        r"(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx])(?!\d)"
    )),
    ("bank_card", re.compile(
        r"(?<!\d)("
        r"[1-9]\d{14,18}"
        r"|"
        r"[1-9]\d{3}(?:[\s\-_.／/]+\d{4}){2,3}(?:[\s\-_.／/]+\d{1,4})?"
        r")(?!\d)"
    )),
    ("phone", re.compile(r"(?<!\d)(1[3-9]\d{9})(?!\d)")),
    ("imei", re.compile(r"(?<!\d)(\d{15})(?!\d)")),
    ("ip", re.compile(r"(?<!\d)((?:\d{1,3}\.){3}\d{1,3})(?!\d)")),
    ("account", re.compile(r"(?i)(?:账号|帐户|账户|user(?:name)?|login)[:：\s]*([A-Za-z0-9_.-]{4,32})")),
    ("address", re.compile(
        r"([\u4e00-\u9fff]{2,10}(?:省|市|自治区|特别行政区))?"
        r"[\u4e00-\u9fff]{1,10}(?:市|州|盟)?"
        r"[\u4e00-\u9fff]{1,12}(?:区|县|旗)"
        r"[\u4e00-\u9fff0-9\-号弄幢栋单元室楼]{0,30}"
    )),
    ("name", re.compile(
        r"(?:"
        r"(?:姓名|被告人|嫌疑人|当事人|原告|被告|证人|辩护人|被害人|申诉人|被申诉人|法定代表人|负责人|联系人)"
        r"[:：\s]*([\u4e00-\u9fff·]{2,4})"
        r"|"
        r"(?<![^\s,，。；：、\(\)（）])([\u4e00-\u9fff·]{2,4})(?![^\s,，。；：、\(\)（）])"
        r")"
    )),
]

_UNREDACTED_PATTERNS = [
    re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)"),
    re.compile(
        r"(?<!\d)[1-9]\d{5}(?:19|20)\d{2}(?:0[1-9]|1[0-2])"
        r"(?:0[1-9]|[12]\d|3[01])\d{3}[\dXx](?!\d)"
    ),
    re.compile(r"(?<!\d)\d{15}(?!\d)"),
]


@dataclass
class RedactionHit:
    sens_type: str
    start: int
    end: int
    original: str
    placeholder: str


def merge_person_spans(text: str, person_results: list) -> list:
    if not person_results:
        return []

    MERGE_SEPARATORS = set("，,、；;。.！!？?：:\"\"''（）()【】[]《》<>／/\\\t\n\r ")
    MERGE_STOP_WORDS = {"的", "和", "与", "及", "或", "以及", "及其", "暨"}
    VERB_PREFIXES = set("受被让给为由将把向对与同跟从在到予以用拿借凭靠沿顺朝往冲离除比")
    VERB_SUFFIXES = ["说", "道", "讲", "问", "答"]

    merged = []
    current = person_results[0]

    for r in person_results[1:]:
        gap = text[current.end : r.start]
        clean_gap = (gap == "") or (
            not any(ch in MERGE_SEPARATORS for ch in gap) and
            not any(w in gap for w in MERGE_STOP_WORDS)
        )
        if clean_gap:
            second_first_char = text[r.start:r.start+1]
            if second_first_char in VERB_PREFIXES:
                merged.append(current)
                current = r
                continue
            new_start = min(current.start, r.start)
            new_end = max(current.end, r.end)
            new_score = max(current.score, r.score)
            current = RecognizerResult(
                entity_type="PERSON",
                start=new_start,
                end=new_end,
                score=new_score
            )
        else:
            merged.append(current)
            current = r
    merged.append(current)

    merged = [r for r in merged if (r.end - r.start) >= 2]

    filtered = []
    for r in merged:
        name = text[r.start:r.end]
        while len(name) >= 2 and name[-1] in VERB_SUFFIXES:
            name = name[:-1]
        if name != text[r.start:r.end]:
            r = RecognizerResult(
                entity_type="PERSON",
                start=r.start,
                end=r.start + len(name),
                score=r.score
            )
        filtered.append(r)

    trimmed = []
    for r in filtered:
        name = text[r.start:r.end]
        trim_count = 0
        while len(name) >= 2 and name[0] in VERB_PREFIXES:
            name = name[1:]
            trim_count += 1
        if trim_count > 0 and len(name) >= 2:
            r = RecognizerResult(
                entity_type="PERSON",
                start=r.start + trim_count,
                end=r.end,
                score=r.score
            )
        trimmed.append(r)
    return trimmed


def mask_phone_number(phone: str) -> str:
    cleaned = phone.replace("+86", "").replace("-", "").replace(" ", "")
    if cleaned.isdigit() and len(cleaned) == 11:
        return cleaned[:3] + "****" + cleaned[-4:]
    return phone


def _generate_placeholder(sens_type: str) -> str:
    short_uuid = uuid.uuid4().hex[:8]
    return f"{sens_type.upper()}_{short_uuid}"


def redact_text(
    text: str,
    document_version_id: str = None,
    chunk_id: str = None,
) -> Tuple[str, List[RedactionHit]]:
    mapper = get_global_mapper()
    analyzer = get_global_analyzer()

    if not text:
        return text, []

    results = analyzer.analyze(
        text=text,
        language="zh",
        entities=[
            "PERSON", "PHONE_NUMBER", "CREDIT_CARD", "ID", "EMAIL_ADDRESS", "URL"
        ],
        score_threshold=0.8
    )

    person_results = [r for r in results if r.entity_type == "PERSON"]
    other_results = [r for r in results if r.entity_type != "PERSON"]

    person_results.sort(key=lambda r: (r.start, r.end))
    person_results = merge_person_spans(text, person_results)
    filtered_person_results = []
    for r in person_results:
        name = text[r.start:r.end]
        # 只允许中文字符和间隔符“·”
        if re.fullmatch(r'[\u4e00-\u9fa5·]+', name):
            filtered_person_results.append(r)
    person_results = filtered_person_results

    spans: list[tuple[int, int, str, str]] = []

    for r in person_results:
        spans.append((r.start, r.end, text[r.start:r.end], "PERSON"))

    VERB_PREFIXES = set("受被让给为由将把向对与同跟从在到予以用拿借凭靠沿顺朝往冲离除比")
    name_spans: dict[str, list[tuple[int, int]]] = {}
    for s, e, name, etype in spans:
        if etype == "PERSON":
            name_spans.setdefault(name, []).append((s, e))

    sorted_names = sorted(name_spans.keys(), key=len, reverse=True)
    alias_map: dict[str, str] = {}
    for short_name in sorted_names:
        if len(short_name) < 2:
            continue
        candidates = [
            ln for ln in sorted_names
            if len(ln) > len(short_name) and ln.startswith(short_name)
        ]
        if not candidates:
            continue
        short_positions = name_spans[short_name]
        found = False
        for long_name in candidates:
            remaining = long_name[len(short_name):]
            if remaining and remaining[0] in VERB_PREFIXES:
                continue
            for sp_start, sp_end in short_positions:
                next_start = sp_end
                if (next_start + len(remaining) <= len(text) and
                        text[next_start:next_start + len(remaining)] == remaining):
                    after_end = next_start + len(remaining)
                    if after_end < len(text):
                        next_char = text[after_end]
                        if '\u4e00' <= next_char <= '\u9fff':
                            continue
                    alias_map[short_name] = long_name
                    found = True
                    break
            if found:
                break

    new_spans = []
    merged_short = set()
    for s, e, name, etype in spans:
        if etype == "PERSON" and name in alias_map:
            merged_short.add(name)
            continue
        new_spans.append((s, e, name, etype))

    for short_name, long_name in alias_map.items():
        for sp_start, sp_end in name_spans[short_name]:
            remaining = long_name[len(short_name):]
            next_start = sp_end
            if (next_start + len(remaining) <= len(text) and
                    text[next_start:next_start + len(remaining)] == remaining):
                new_spans.append((sp_start, next_start + len(remaining), long_name, "PERSON"))

    seen = set()
    unique_spans = []
    for span in new_spans:
        key = (span[0], span[1], span[2], span[3])
        if key not in seen:
            seen.add(key)
            unique_spans.append(span)
    spans = unique_spans

    for r in other_results:
        spans.append((r.start, r.end, text[r.start:r.end], r.entity_type))

    spans.sort(key=lambda x: (x[0], -(x[1] - x[0])))
    merged_spans: list[tuple[int, int, str, str]] = []
    last_end = -1
    for start, end, original, entity_type in spans:
        if start < last_end:
            continue
        merged_spans.append((start, end, original, entity_type))
        last_end = end

    redacted_text = text
    hits = []
    for start, end, original, entity_type in reversed(merged_spans):
        if entity_type == "PERSON":
            placeholder = mapper.get_or_create(original, "PERSON")
        elif entity_type == "PHONE_NUMBER":
            masked = mask_phone_number(original)
            placeholder = masked if masked != original else mapper.get_or_create(original, "phone")
        else:
            placeholder = mapper.get_or_create(original, entity_type.lower())

        redacted_text = redacted_text[:start] + placeholder + redacted_text[end:]
        hits.append(RedactionHit(
            sens_type=entity_type,
            start=start,
            end=end,
            original=original,
            placeholder=placeholder
        ))

    hits.sort(key=lambda h: h.start)
    return redacted_text, hits


# ---------- 解析与页级 OCR ----------
@dataclass
class PageResult:
    page_no: int
    source: str
    text: str
    text_density: float = 0.0
    avg_confidence: float | None = None
    min_confidence: float | None = None
    bbox: list[Any] = field(default_factory=list)
    lines: list[dict[str, Any]] = field(default_factory=list)
    quality_flags: list[str] = field(default_factory=list)
    status: str = "PARSED"
    error_code: str | None = None
    error_message: str | None = None
    image_bytes: bytes | None = None


@dataclass
class OCRLine:
    text: str
    confidence: float
    bbox: list[float]


class OCREngine(Protocol):
    def recognize(self, image_bytes: bytes) -> list[OCRLine]:
        ...


class FallbackOCREngine:
    def recognize(self, image_bytes: bytes) -> list[OCRLine]:
        marker = b"__OCR_TEXT__:"
        if marker in image_bytes:
            text = image_bytes.split(marker, 1)[1].decode("utf-8", errors="ignore")
            return [OCRLine(text, 0.95 if text.strip() else 0.2, [0, 0, 100, 20])]

        printable = "".join(chr(value) if 32 <= value < 127 else " " for value in image_bytes)
        tokens = [token for token in printable.split() if len(token) >= 4]
        if not tokens:
            return [OCRLine("", 0.1, [0, 0, 1, 1])]
        return [OCRLine(" ".join(tokens[:50]), 0.55, [0, 0, 100, 20])]


class PaddleOCREngine:
    def __init__(self):
        from paddleocr import PaddleOCR
        self._ocr = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)

    def recognize(self, image_bytes: bytes) -> list[OCRLine]:
        import numpy as np
        from PIL import Image

        image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
        lines: list[OCRLine] = []
        for block in self._ocr.ocr(np.array(image), cls=True) or []:
            for box, (text, confidence) in block or []:
                flat = [float(value) for point in box for value in point]
                xs, ys = flat[0::2], flat[1::2]
                lines.append(
                    OCRLine(str(text), float(confidence), [min(xs), min(ys), max(xs), max(ys)])
                )
        return lines


_ocr_engine: OCREngine | None = None


def set_ocr_engine(engine: OCREngine | None) -> None:
    global _ocr_engine
    _ocr_engine = engine


def _get_ocr_engine() -> OCREngine:
    global _ocr_engine
    if _ocr_engine is None:
        try:
            _ocr_engine = PaddleOCREngine()
        except Exception:
            _ocr_engine = FallbackOCREngine()
    return _ocr_engine


def _recognize_image(image_bytes: bytes) -> dict[str, Any]:
    lines = _get_ocr_engine().recognize(image_bytes)
    texts = [line.text for line in lines if line.text]
    confidences = [line.confidence for line in lines] or [0.0]
    average = sum(confidences) / len(confidences)
    minimum = min(confidences)
    flags = []
    if minimum < OCR_LOW_CONFIDENCE_THRESHOLD or average < OCR_LOW_CONFIDENCE_THRESHOLD:
        flags.append("LOW_OCR_CONFIDENCE")
    if not "".join(texts).strip():
        flags.append("EMPTY_OCR")
    return {
        "text": "\n".join(texts),
        "lines": [
            {"text": line.text, "confidence": line.confidence, "bbox": line.bbox}
            for line in lines
        ],
        "avg_confidence": average,
        "min_confidence": minimum,
        "quality_flags": flags,
        "bbox": [line.bbox for line in lines],
    }


def _extract_pdf(path: Path) -> list[PageResult]:
    try:
        import fitz
    except ImportError as exc:
        raise MaterialError(ERROR_CODES["PARSE_FAILED"], "PyMuPDF(fitz) not installed") from exc

    try:
        document = fitz.open(path)
    except Exception as exc:
        message = str(exc).lower()
        code = (
            ERROR_CODES["ENCRYPTED_FILE"]
            if "password" in message or "encrypt" in message
            else ERROR_CODES["CORRUPT_FILE"]
        )
        raise MaterialError(code, f"Failed to open PDF: {exc}") from exc

    needs_password = getattr(document, "needs_pass", False) or getattr(
        document, "is_encrypted", False
    )
    if needs_password and not document.authenticate(""):
        document.close()
        raise MaterialError(ERROR_CODES["ENCRYPTED_FILE"], "PDF requires a password")

    pages: list[PageResult] = []
    try:
        for index, page in enumerate(document):
            try:
                text = page.get_text("text") or ""
                bbox = [
                    [float(block[0]), float(block[1]), float(block[2]), float(block[3])]
                    for block in (page.get_text("blocks") or [])
                    if len(block) >= 4
                ]
                area = float(page.rect.width * page.rect.height) if page.rect else 1.0
                density = len(text.strip()) / max(area / 10000.0, 1.0)
                needs_ocr = density < OCR_TEXT_DENSITY_THRESHOLD or not text.strip()
                pages.append(
                    PageResult(
                        page_no=index + 1,
                        source="pdf_text",
                        text=text,
                        text_density=density,
                        bbox=bbox,
                        status="NEEDS_OCR" if needs_ocr else "PARSED",
                        image_bytes=(
                            page.get_pixmap(matrix=fitz.Matrix(2, 2)).tobytes("png")
                            if needs_ocr
                            else None
                        ),
                        quality_flags=["LOW_TEXT_DENSITY"] if needs_ocr else [],
                    )
                )
            except Exception as exc:
                pages.append(
                    PageResult(
                        page_no=index + 1,
                        source="pdf_text",
                        text="",
                        status="FAILED",
                        error_code=ERROR_CODES["PARSE_FAILED"],
                        error_message=str(exc),
                        quality_flags=["PAGE_PARSE_FAILED"],
                    )
                )
    finally:
        document.close()
    return pages


def _extract_docx(path: Path) -> list[PageResult]:
    try:
        from docx import Document
        document = Document(str(path))
    except ImportError as exc:
        raise MaterialError(ERROR_CODES["PARSE_FAILED"], "python-docx not installed") from exc
    except Exception as exc:
        raise MaterialError(ERROR_CODES["CORRUPT_FILE"], f"Cannot open DOCX: {exc}") from exc

    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return [PageResult(1, "docx", "\n".join(parts), text_density=1.0)]


def _extract_txt(path: Path) -> list[PageResult]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            text = path.read_text(encoding="gbk")
        except Exception as exc:
            raise MaterialError(
                ERROR_CODES["CORRUPT_FILE"], f"Cannot decode text file: {exc}"
            ) from exc
    except Exception as exc:
        raise MaterialError(ERROR_CODES["CORRUPT_FILE"], f"Cannot read text file: {exc}") from exc
    return [PageResult(1, "txt", text, text_density=1.0)]


def _extract_image(path: Path) -> list[PageResult]:
    return [
        PageResult(
            1,
            "image",
            "",
            status="NEEDS_OCR",
            image_bytes=path.read_bytes(),
            quality_flags=["IMAGE_REQUIRES_OCR"],
        )
    ]


def _run_page_ocr(page: PageResult, retries: int) -> PageResult:
    if not page.image_bytes:
        page.status = "OCR_FAILED"
        page.error_code = ERROR_CODES["OCR_FAILED"]
        page.error_message = "missing page image for OCR"
        page.quality_flags = sorted(set(page.quality_flags + ["OCR_FAILED"]))
        return page

    last_error: Exception | None = None
    for _ in range(retries + 1):
        try:
            result = _recognize_image(page.image_bytes)
            page.text = result["text"]
            page.lines = result["lines"]
            page.bbox = result["bbox"]
            page.avg_confidence = result["avg_confidence"]
            page.min_confidence = result["min_confidence"]
            page.quality_flags = sorted(set(page.quality_flags + result["quality_flags"]))
            page.source = "ocr"
            if "EMPTY_OCR" in result["quality_flags"]:
                page.status = "OCR_FAILED"
                page.error_code = ERROR_CODES["OCR_FAILED"]
                page.error_message = "OCR returned empty text"
            elif "LOW_OCR_CONFIDENCE" in result["quality_flags"]:
                page.status = "NEEDS_OCR_REVIEW"
            else:
                page.status = "PARSED"
            return page
        except Exception as exc:
            last_error = exc

    page.status = "OCR_FAILED"
    page.error_code = ERROR_CODES["OCR_FAILED"]
    page.error_message = str(last_error) if last_error else "OCR failed"
    page.quality_flags = sorted(set(page.quality_flags + ["OCR_FAILED"]))
    return page


def parse_file_to_pages(path: Path | str, max_retries: int | None = None) -> list[dict[str, Any]]:
    file_path = Path(path)
    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".txt": _extract_txt,
        ".png": _extract_image,
        ".jpg": _extract_image,
        ".jpeg": _extract_image,
    }
    extractor = extractors.get(file_path.suffix.lower())
    if extractor is None:
        raise MaterialError(
            ERROR_CODES["UNSUPPORTED_TYPE"], f"Unsupported extension: {file_path.suffix.lower()}"
        )

    retries = OCR_MAX_PAGE_RETRIES if max_retries is None else max_retries
    parsed = []
    for page in extractor(file_path):
        if page.status == "NEEDS_OCR":
            page = _run_page_ocr(page, retries)
        page.image_bytes = None
        parsed.append(
            {
                "page_no": page.page_no,
                "source": page.source,
                "text": page.text,
                "text_density": page.text_density,
                "avg_confidence": page.avg_confidence,
                "min_confidence": page.min_confidence,
                "bbox": page.bbox,
                "lines": page.lines,
                "quality_flags": page.quality_flags,
                "status": page.status,
                "error_code": page.error_code,
                "error_message": page.error_message,
            }
        )
    return parsed


def summarize_page_quality(pages: list[dict[str, Any]]) -> dict[str, Any]:
    confidences = [p["avg_confidence"] for p in pages if p.get("avg_confidence") is not None]
    minimums = [p["min_confidence"] for p in pages if p.get("min_confidence") is not None]
    low_pages: list[int] = []
    warnings: list[str] = []
    statuses: dict[str, str] = {}
    for page in pages:
        page_no = str(page["page_no"])
        status = page.get("status") or "PARSED"
        flags = page.get("quality_flags") or []
        statuses[page_no] = status
        if status in {"OCR_FAILED", "FAILED"}:
            warnings.append(f"page {page_no}: {status}")
        if "LOW_OCR_CONFIDENCE" in flags:
            low_pages.append(int(page["page_no"]))
            warnings.append(f"page {page_no}: low OCR confidence")
        if "EMPTY_OCR" in flags:
            warnings.append(f"page {page_no}: empty OCR")
    return {
        "avg_confidence": sum(confidences) / len(confidences) if confidences else None,
        "min_confidence": min(minimums) if minimums else None,
        "low_confidence_pages": sorted(set(low_pages)),
        "warnings": warnings,
        "page_statuses": statuses,
    }


def derive_version_status(pages: list[dict[str, Any]]) -> str:
    statuses = [page.get("status") for page in pages]
    if not statuses:
        return "FAILED"
    if any(status == "OCR_FAILED" for status in statuses):
        return "OCR_FAILED"
    if any(status == "NEEDS_OCR_REVIEW" for status in statuses):
        return "NEEDS_OCR_REVIEW"
    if any(status == "FAILED" for status in statuses):
        return "FAILED"
    return "PARSED" if all(status == "PARSED" for status in statuses) else "NEEDS_OCR_REVIEW"


# ---------- 分块 ----------
def _text_sha256(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def build_chunks(
    document_version_id: str,
    pages: list[dict[str, Any]],
    *,
    chunk_size: int | None = None,
    overlap: int | None = None,
    parser_version: str | None = None,
) -> list[dict[str, Any]]:
    size = CHUNK_SIZE if chunk_size is None else chunk_size
    overlap_size = CHUNK_OVERLAP if overlap is None else overlap
    version = parser_version or PARSER_VERSION
    if size <= 0:
        raise ValueError("chunk_size must be positive")
    if overlap_size >= size:
        overlap_size = max(0, size // 4)

    stream: list[tuple[str, int, list]] = []
    for page in pages:
        page_no = int(page["page_no"])
        text = page.get("text") or ""
        bbox = page.get("bbox") or []
        stream.extend((character, page_no, bbox) for character in text)
        if text and not text.endswith("\n"):
            stream.append(("\n", page_no, bbox))

    full_text = "".join(character for character, _, _ in stream)
    if not full_text.strip():
        digest = _text_sha256("")
        return [
            {
                "id": f"{document_version_id}:0:{digest[:16]}",
                "document_version_id": document_version_id,
                "ordinal": 0,
                "page_start": int(pages[0]["page_no"]) if pages else 1,
                "page_end": int(pages[-1]["page_no"]) if pages else 1,
                "char_start": 0,
                "char_end": 0,
                "bbox_json": "[]",
                "text_raw": "",
                "text_redacted": "",
                "text_sha256": digest,
                "parser_version": version,
                "quality_flags_json": "[]",
                "is_active": 1,
                "stale": 0,
            }
        ]

    chunks: list[dict[str, Any]] = []
    start = 0
    ordinal = 0
    while start < len(full_text):
        end = min(start + size, len(full_text))
        if end < len(full_text):
            window = full_text[start:end]
            break_at = max(window.rfind("\n"), window.rfind(" "), window.rfind("\u3000"))
            if break_at > size * 0.5:
                end = start + break_at + 1

        text = full_text[start:end]
        page_start = stream[start][1]
        page_end = stream[end - 1][1]
        flags: list[str] = []
        for page in pages:
            if page_start <= int(page["page_no"]) <= page_end:
                flags.extend(page.get("quality_flags") or [])
        digest = _text_sha256(text)
        chunks.append(
            {
                "id": f"{document_version_id}:{ordinal}:{digest[:16]}",
                "document_version_id": document_version_id,
                "ordinal": ordinal,
                "page_start": page_start,
                "page_end": page_end,
                "char_start": start,
                "char_end": end,
                "bbox_json": json.dumps(stream[start][2], ensure_ascii=False),
                "text_raw": text,
                "text_redacted": text,
                "text_sha256": digest,
                "parser_version": version,
                "quality_flags_json": json.dumps(sorted(set(flags)), ensure_ascii=False),
                "is_active": 1,
                "stale": 0,
            }
        )
        ordinal += 1
        if end >= len(full_text):
            break
        start = max(end - overlap_size, start + 1)
    return chunks


# ---------- MaterialService ----------
class MaterialService:
    def __init__(self, db_path=None, auth_check=None, mapper_salt=None):
        self.db_path = db_path
        self.auth_check = auth_check or _default_auth()
        init_db(self.db_path)

    def _init_analyzer(self) -> AnalyzerEngine | None:
        try:
            from presidio_analyzer import AnalyzerEngine, PatternRecognizer, Pattern
            from presidio_analyzer.nlp_engine import NlpEngineProvider

            configuration = {
                "nlp_engine_name": "spacy",
                "models": [{"lang_code": "zh", "model_name": "zh_core_web_trf"}],
            }
            provider = NlpEngineProvider(nlp_configuration=configuration)
            nlp_engine = provider.create_engine()
            analyzer = AnalyzerEngine(nlp_engine=nlp_engine)

            chinese_name_pattern = Pattern(
                name="chinese_name_pattern",
                regex=r"(?:姓名|被告人|嫌疑人|当事人|原告|被告|证人|辩护人|被害人|申诉人|被申诉人|法定代表人|负责人|联系人|申请人|被申请人)[:：\s]*([\u4e00-\u9fa5·]{2,4})",
                score=0.9,
            )
            chinese_name_recognizer = PatternRecognizer(
                supported_entity="PERSON",
                name="chinese_name_recognizer",
                patterns=[chinese_name_pattern],
                context=["姓名", "被告人", "原告", "被告", "证人", "受害人", "当事人"],
            )
            analyzer.registry.add_recognizer(chinese_name_recognizer)

            bare_name_pattern = Pattern(
                name="bare_chinese_name_pattern",
                regex=r"(?<![\u4e00-\u9fa5])([\u4e00-\u9fa5]{2,3})(?![\u4e00-\u9fa5])",
                score=0.4,
            )
            bare_name_recognizer = PatternRecognizer(
                supported_entity="PERSON",
                name="bare_chinese_name_recognizer",
                patterns=[bare_name_pattern],
            )
            analyzer.registry.add_recognizer(bare_name_recognizer)

            return analyzer
        except Exception as e:
            print(f"⚠️ AnalyzerEngine 初始化失败，将使用正则降级: {e}")
            return None

    def _authorize(self, user_id: str | None, case_id: str | None, action: str) -> None:
        allowed, reason = self.auth_check(user_id, case_id, action)
        if not allowed:
            with db_session(self.db_path) as conn:
                add_audit(
                    conn,
                    event_type="auth_check",
                    actor_user_id=user_id,
                    case_id=case_id,
                    decision="deny",
                    reason=reason or "denied",
                    detail_json=json.dumps({"action": action}, ensure_ascii=False),
                )
            raise MaterialError(ERROR_CODES["AUTH_DENIED"], reason or "authorization denied")

    # ----- 上传与版本 -----
    def upload_many(
        self,
        items: list[dict[str, Any]],
        *,
        user_id: str | None = None,
        parse: bool = True,
        keep_duplicate: bool = False,
    ) -> list[dict[str, Any]]:
        return [
            self.upload_one(
                case_id=item["case_id"],
                filename=item["filename"],
                content=item.get("content"),
                path=item.get("path"),
                user_id=user_id,
                parse=parse,
                keep_duplicate=keep_duplicate,
                replace_document_id=item.get("replace_document_id"),
            )
            for item in items
        ]

    def upload_one(
        self,
        *,
        case_id: str,
        filename: str,
        content: bytes | None = None,
        path: str | Path | None = None,
        user_id: str | None = None,
        parse: bool = True,
        keep_duplicate: bool = False,
        replace_document_id: str | None = None,
    ) -> dict[str, Any]:
        self._authorize(user_id, case_id, "material.upload")
        safe_name = Path(filename).name
        extension = Path(safe_name).suffix.lower()
        if extension not in ALLOWED_MATERIAL_EXTENSIONS:
            raise MaterialError(ERROR_CODES["UNSUPPORTED_TYPE"], f"unsupported type: {extension}")

        if content is None:
            if path is None:
                raise MaterialError(ERROR_CODES["CORRUPT_FILE"], "empty upload")
            content = Path(path).read_bytes()

        size = len(content)
        if size > MAX_UPLOAD_BYTES:
            raise MaterialError(
                ERROR_CODES["FILE_TOO_LARGE"],
                f"file exceeds limit {MAX_UPLOAD_BYTES} bytes",
                {"size": size},
            )
        if size == 0:
            raise MaterialError(ERROR_CODES["CORRUPT_FILE"], "empty file")

        digest = hashlib.sha256(content).hexdigest()
        content_type = mimetypes.guess_type(safe_name)[0] or "application/octet-stream"

        with db_session(self.db_path) as conn:
            if not get_case(conn, case_id):
                if user_id == "system":
                    ensure_demo_case(conn, case_id)
                else:
                    raise MaterialError(ERROR_CODES["NOT_FOUND"], f"case not found: {case_id}")

            duplicate = _row(
                conn,
                """
                SELECT d.* FROM documents d
                JOIN document_versions v ON v.document_id = d.id
                WHERE d.case_id = ? AND v.sha256 = ? AND d.deleted_at IS NULL
                ORDER BY v.created_at DESC LIMIT 1
                """,
                (case_id, digest),
            )
            if duplicate and not replace_document_id and not keep_duplicate:
                _update(conn, "documents", duplicate["id"], {"status": "DUPLICATE_PENDING"})
                add_audit(
                    conn,
                    event_type="duplicate_pending",
                    actor_user_id=user_id,
                    case_id=case_id,
                    document_id=duplicate["id"],
                    decision="pending",
                    reason="same sha256 in case",
                    detail_json=json.dumps({"sha256": digest}, ensure_ascii=False),
                )
                return {
                    "status": "DUPLICATE_PENDING",
                    "error_code": ERROR_CODES["DUPLICATE_PENDING"],
                    "document": get_document(conn, duplicate["id"]),
                    "message": "duplicate hash in case; decide keep or cancel",
                }

            now = utc_now()
            if replace_document_id:
                document = get_document(conn, replace_document_id)
                if not document or document["case_id"] != case_id:
                    raise MaterialError(ERROR_CODES["NOT_FOUND"], "document to replace not found")
                document_id = document["id"]
                parent_version_id = document.get("current_version_id")
                source_type = "REPLACE"
                if parent_version_id:
                    deactivate_version_chunks(conn, parent_version_id)
                    _update(conn, "document_versions", parent_version_id, {"is_current": 0})
            else:
                document_id = new_id()
                parent_version_id = None
                source_type = "UPLOAD"
                _insert(
                    conn,
                    "documents",
                    {
                        "id": document_id,
                        "case_id": case_id,
                        "filename": safe_name,
                        "stored_name": safe_name,
                        "content_type": content_type,
                        "size": size,
                        "sha256": digest,
                        "uploaded_by": user_id or "anonymous",
                        "created_at": now,
                        "status": "UPLOADED",
                        "current_version_id": None,
                        "deleted_at": None,
                        "quality_summary_json": "{}",
                    },
                )

            version_id = new_id()
            version_no = next_version_no(conn, document_id)
            storage_path = (
                Path(MATERIAL_STORAGE_DIR)
                / case_id
                / document_id
                / f"v{version_no}_{digest[:12]}{extension}"
            )
            storage_path.parent.mkdir(parents=True, exist_ok=True)
            storage_path.write_bytes(content)

            _insert(
                conn,
                "document_versions",
                {
                    "id": version_id,
                    "document_id": document_id,
                    "version_no": version_no,
                    "parent_version_id": parent_version_id,
                    "source_type": source_type,
                    "sha256": digest,
                    "storage_path": str(storage_path),
                    "content_type": content_type,
                    "size": size,
                    "parser_version": PARSER_VERSION,
                    "status": "UPLOADED",
                    "is_current": 1,
                    "is_active": 1,
                    "quality_summary_json": "{}",
                    "error_code": None,
                    "error_message": None,
                    "created_by": user_id or "anonymous",
                    "created_at": now,
                },
            )
            set_current_version(conn, document_id, version_id)
            _update(
                conn,
                "documents",
                document_id,
                {
                    "filename": safe_name,
                    "stored_name": storage_path.name,
                    "content_type": content_type,
                    "size": size,
                    "sha256": digest,
                    "status": "UPLOADED",
                },
            )
            add_audit(
                conn,
                event_type="upload",
                actor_user_id=user_id,
                case_id=case_id,
                document_id=document_id,
                document_version_id=version_id,
                decision="allow",
                reason=source_type.lower(),
            )

        if parse:
            return self.parse_version(version_id, user_id=user_id)
        with db_session(self.db_path) as conn:
            return {
                "document": get_document(conn, document_id),
                "version": get_version(conn, version_id),
                "status": "UPLOADED",
            }

    # ----- 解析 -----
    def parse_version(self, version_id: str, user_id: str | None = None) -> dict[str, Any]:
        with db_session(self.db_path) as conn:
            version = get_version(conn, version_id)
            if not version:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "version not found")
            document = get_document(conn, version["document_id"])
            case_id = document["case_id"] if document else None
            storage_path = Path(version["storage_path"])

        self._authorize(user_id, case_id, "material.parse")

        job_id = new_id()
        with db_session(self.db_path) as conn:
            now = utc_now()
            _insert(
                conn,
                "parse_jobs",
                {
                    "id": job_id,
                    "document_version_id": version_id,
                    "status": "RUNNING",
                    "attempt": 1,
                    "max_attempts": OCR_MAX_PAGE_RETRIES,
                    "error_code": None,
                    "error_message": None,
                    "started_at": now,
                    "finished_at": None,
                    "created_at": now,
                },
            )
            _update(conn, "document_versions", version_id, {"status": "PARSING"})
            if document:
                _update(conn, "documents", document["id"], {"status": "PARSING"})

        try:
            pages = parse_file_to_pages(storage_path)
        except Exception as exc:
            code = exc.code if isinstance(exc, MaterialError) else ERROR_CODES["PARSE_FAILED"]
            message = exc.message if isinstance(exc, MaterialError) else str(exc)
            with db_session(self.db_path) as conn:
                _update(
                    conn,
                    "parse_jobs",
                    job_id,
                    {
                        "status": "FAILED",
                        "error_code": code,
                        "error_message": message,
                        "finished_at": utc_now(),
                    },
                )
                _update(
                    conn,
                    "document_versions",
                    version_id,
                    {"status": "FAILED", "error_code": code, "error_message": message},
                )
                if document:
                    _update(conn, "documents", document["id"], {"status": "FAILED"})
            if isinstance(exc, MaterialError):
                raise
            raise MaterialError(code, message) from exc

        return self._persist_pages(
            version_id=version_id,
            document_id=version["document_id"],
            case_id=case_id,
            pages=pages,
            user_id=user_id,
            job_id=job_id,
            event_type="parse",
        )

    def _persist_pages(
        self,
        *,
        version_id: str,
        document_id: str,
        case_id: str | None,
        pages: list[dict[str, Any]],
        user_id: str | None,
        job_id: str | None,
        event_type: str,
        audit_reason: str | None = None,
    ) -> dict[str, Any]:
        quality = summarize_page_quality(pages)
        status = derive_version_status(pages)
        chunks = build_chunks(version_id, pages)

        redaction_rows: list[dict[str, Any]] = []
        for chunk in chunks:
            redacted, hits = redact_text(chunk["text_raw"])
            chunk["text_redacted"] = redacted
            chunk["_hits"] = hits
            for hit in hits:
                redaction_rows.append({
                    "id": new_id(),
                    "document_version_id": version_id,
                    "chunk_id": chunk["id"],
                    "sens_type": hit.sens_type,
                    "start_offset": hit.start,
                    "end_offset": hit.end,
                    "placeholder": hit.placeholder,
                    "map_ref": get_global_mapper()._fingerprint(hit.original),
                    "created_at": utc_now(),
                })

        with db_session(self.db_path) as conn:
            for page in pages:
                upsert_page(
                    conn,
                    {
                        "id": new_id(),
                        "document_version_id": version_id,
                        "page_no": page["page_no"],
                        "source": page["source"],
                        "text": page["text"],
                        "text_density": page.get("text_density") or 0,
                        "avg_confidence": page.get("avg_confidence"),
                        "min_confidence": page.get("min_confidence"),
                        "bbox_json": json.dumps(page.get("bbox") or [], ensure_ascii=False),
                        "lines_json": json.dumps(page.get("lines") or [], ensure_ascii=False),
                        "quality_flags_json": json.dumps(
                            page.get("quality_flags") or [], ensure_ascii=False
                        ),
                        "status": page.get("status") or "PARSED",
                        "retry_count": 0,
                        "error_code": page.get("error_code"),
                        "error_message": page.get("error_message"),
                    },
                )

            for chunk in chunks:
                chunk.pop("_hits", None)
            replace_chunks(conn, version_id, chunks)

            for row in redaction_rows:
                _insert(conn, "redaction_items", row)

            quality_json = json.dumps(quality, ensure_ascii=False)
            _update(
                conn,
                "document_versions",
                version_id,
                {"status": status, "quality_summary_json": quality_json},
            )
            _update(
                conn,
                "documents",
                document_id,
                {"status": status, "quality_summary_json": quality_json},
            )
            if job_id:
                _update(conn, "parse_jobs", job_id, {"status": "DONE", "finished_at": utc_now()})
            add_audit(
                conn,
                event_type=event_type,
                actor_user_id=user_id,
                case_id=case_id,
                document_id=document_id,
                document_version_id=version_id,
                decision="allow",
                reason=audit_reason or status,
            )
            return {
                "document": get_document(conn, document_id),
                "version": get_version(conn, version_id),
                "pages": list_pages(conn, version_id),
                "chunks": [
                    {key: value for key, value in chunk.items() if key != "text_raw"}
                    for chunk in list_chunks(conn, version_id)
                ],
                "quality": quality,
                "status": status,
            }

    # ----- 查询 -----
    def list_materials(self, case_id: str, user_id: str | None = None) -> list[dict[str, Any]]:
        self._authorize(user_id, case_id, "material.list")
        with db_session(self.db_path) as conn:
            documents = _rows(
                conn,
                "SELECT * FROM documents WHERE case_id = ? AND deleted_at IS NULL ORDER BY created_at DESC",
                (case_id,),
            )
            materials = []
            for document in documents:
                versions = list_versions(conn, document["id"])
                materials.append(
                    {
                        **document,
                        "quality_summary": json.loads(document.get("quality_summary_json") or "{}"),
                        "current_version": next(
                            (v for v in versions if v.get("is_current")), None
                        ),
                        "version_count": len(versions),
                    }
                )
            return materials

    def get_status(self, document_id: str, user_id: str | None = None) -> dict[str, Any]:
        with db_session(self.db_path) as conn:
            document = get_document(conn, document_id)
            if not document:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "document not found")
            case_id = document["case_id"]

        self._authorize(user_id, case_id, "material.read")

        with db_session(self.db_path) as conn:
            document = get_document(conn, document_id)
            current = (
                get_version(conn, document["current_version_id"])
                if document.get("current_version_id")
                else None
            )
            pages = list_pages(conn, current["id"]) if current else []
            low_quality = []
            for page in pages:
                flags = json.loads(page.get("quality_flags_json") or "[]")
                if page["status"] in {"NEEDS_OCR_REVIEW", "OCR_FAILED", "FAILED"} or (
                    "LOW_OCR_CONFIDENCE" in flags
                ):
                    low_quality.append(
                        {
                            "page_no": page["page_no"],
                            "status": page["status"],
                            "quality_flags": flags,
                            "min_confidence": page.get("min_confidence"),
                            "avg_confidence": page.get("avg_confidence"),
                        }
                    )
            return {
                "document": document,
                "versions": list_versions(conn, document_id),
                "current_version": current,
                "quality_summary": json.loads(document.get("quality_summary_json") or "{}"),
                "low_quality_pages": low_quality,
            }

    # ----- 外发门控 -----
    def read_redacted_chunk(
        self,
        document_version_id: str,
        chunk_id: str | None = None,
        user_id: str | None = None,
    ) -> dict[str, Any]:
        with db_session(self.db_path) as conn:
            version = get_version(conn, document_version_id)
            if not version:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "version not found")
            document = get_document(conn, version["document_id"])
            case_id = document["case_id"] if document else None

        self._authorize(user_id, case_id, "material.read_redacted")

        with db_session(self.db_path) as conn:
            version = get_version(conn, document_version_id)
            document_id = version["document_id"]

            def deny(reason: str, message: str, code: str, detail: dict | None = None):
                add_audit(
                    conn,
                    event_type="egress_check",
                    actor_user_id=user_id,
                    case_id=case_id,
                    document_id=document_id,
                    document_version_id=document_version_id,
                    decision="deny",
                    reason=reason,
                    detail_json=json.dumps(detail or {}, ensure_ascii=False),
                )
                return MaterialError(code, message)

            if not version.get("is_active"):
                raise deny(
                    "version_inactive",
                    "inactive version cannot egress",
                    ERROR_CODES["EGRESS_DENIED"],
                )

            if chunk_id:
                chunk = _row(conn, "SELECT * FROM document_chunks WHERE id = ?", (chunk_id,))
            else:
                available = list_chunks(conn, document_version_id, active_only=True)
                chunk = available[0] if available else None
            if not chunk:
                raise deny("chunk_not_found", "chunk not found", ERROR_CODES["NOT_FOUND"])

            if not chunk.get("is_active") or chunk.get("stale"):
                raise deny(
                    "chunk_stale_or_inactive",
                    "stale/inactive chunk cannot egress",
                    ERROR_CODES["EGRESS_DENIED"],
                    {"chunk_id": chunk["id"]},
                )

            redactions = _rows(
                conn,
                "SELECT * FROM redaction_items WHERE document_version_id = ?",
                (document_version_id,),
            )
            unchanged = chunk.get("text_raw") and chunk["text_redacted"] == chunk["text_raw"]
            if unchanged and redactions:
                raise deny(
                    "unredacted_with_items",
                    "unredacted text cannot egress",
                    ERROR_CODES["EGRESS_DENIED"],
                    {"chunk_id": chunk["id"]},
                )
            text = chunk.get("text_redacted") or ""
            if any(pattern.search(text) for pattern in _UNREDACTED_PATTERNS):
                raise deny(
                    "looks_unredacted",
                    "unredacted sensitive text cannot egress",
                    ERROR_CODES["EGRESS_DENIED"],
                    {"chunk_id": chunk["id"]},
                )

            add_audit(
                conn,
                event_type="egress_check",
                actor_user_id=user_id,
                case_id=case_id,
                document_id=document_id,
                document_version_id=document_version_id,
                decision="allow",
                reason="ok",
                detail_json=json.dumps({"chunk_id": chunk["id"]}, ensure_ascii=False),
            )
            return {
                "chunk_id": chunk["id"],
                "document_version_id": document_version_id,
                "ordinal": chunk["ordinal"],
                "page_start": chunk["page_start"],
                "page_end": chunk["page_end"],
                "text": chunk["text_redacted"],
                "quality_flags": json.loads(chunk.get("quality_flags_json") or "[]"),
                "redacted": True,
            }

    # ----- 人工修正 -----
    def apply_correction(
        self,
        *,
        document_id: str,
        source_version_id: str,
        page_no: int,
        corrected_text: str,
        user_id: str | None = None,
    ) -> dict[str, Any]:
        with db_session(self.db_path) as conn:
            document = get_document(conn, document_id)
            if not document:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "document not found")
            case_id = document["case_id"]

        self._authorize(user_id, case_id, "material.correct")

        with db_session(self.db_path) as conn:
            source = get_version(conn, source_version_id)
            if not source or source["document_id"] != document_id:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "source version not found")
            pages = list_pages(conn, source_version_id)
            if not pages:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "source pages not found")

            corrected_pages: list[dict[str, Any]] = []
            found = False
            for page in pages:
                item = {
                    "page_no": page["page_no"],
                    "source": page["source"],
                    "text": page["text"],
                    "text_density": page["text_density"],
                    "avg_confidence": page["avg_confidence"],
                    "min_confidence": page["min_confidence"],
                    "bbox": json.loads(page.get("bbox_json") or "[]"),
                    "lines": json.loads(page.get("lines_json") or "[]"),
                    "quality_flags": json.loads(page.get("quality_flags_json") or "[]"),
                    "status": "PARSED",
                    "error_code": None,
                    "error_message": None,
                }
                if int(page["page_no"]) == int(page_no):
                    found = True
                    item.update(
                        {
                            "text": corrected_text,
                            "source": "human_correction",
                            "avg_confidence": 1.0,
                            "min_confidence": 1.0,
                            "quality_flags": ["HUMAN_CORRECTED"],
                        }
                    )
                corrected_pages.append(item)
            if not found:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], f"page {page_no} not found")

            version_id = new_id()
            deactivate_version_chunks(conn, source_version_id)
            _update(conn, "document_versions", source_version_id, {"is_current": 0})
            _insert(
                conn,
                "document_versions",
                {
                    "id": version_id,
                    "document_id": document_id,
                    "version_no": next_version_no(conn, document_id),
                    "parent_version_id": source_version_id,
                    "source_type": "CORRECTION",
                    "sha256": source["sha256"],
                    "storage_path": source["storage_path"],
                    "content_type": source["content_type"],
                    "size": source["size"],
                    "parser_version": PARSER_VERSION,
                    "status": "PARSING",
                    "is_current": 1,
                    "is_active": 1,
                    "quality_summary_json": "{}",
                    "error_code": None,
                    "error_message": None,
                    "created_by": user_id or "anonymous",
                    "created_at": utc_now(),
                },
            )
            set_current_version(conn, document_id, version_id)

        result = self._persist_pages(
            version_id=version_id,
            document_id=document_id,
            case_id=case_id,
            pages=corrected_pages,
            user_id=user_id,
            job_id=None,
            event_type="correction",
            audit_reason=f"correct_page_{page_no}",
        )
        result["parent_version_id"] = source_version_id
        return result

    # ----- 删除与重复裁决 -----
    def preview_delete_impact(self, document_id: str, user_id: str | None = None) -> dict[str, Any]:
        with db_session(self.db_path) as conn:
            document = get_document(conn, document_id)
            if not document:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "document not found")
            case_id = document["case_id"]

        self._authorize(user_id, case_id, "material.delete")

        with db_session(self.db_path) as conn:
            impact = []
            for version in list_versions(conn, document_id):
                active = _row(
                    conn,
                    "SELECT COUNT(*) AS total FROM document_chunks WHERE document_version_id = ? AND is_active = 1",
                    (version["id"],),
                )["total"]
                impact.append(
                    {
                        "kind": "version",
                        "id": version["id"],
                        "detail": f"v{version['version_no']} status={version['status']} active_chunks={active}",
                    }
                )
                for chunk in list_chunks(conn, version["id"], active_only=False):
                    impact.append(
                        {
                            "kind": "chunk",
                            "id": chunk["id"],
                            "detail": f"ordinal={chunk['ordinal']} stale={chunk['stale']}",
                        }
                    )
            return {"document_id": document_id, "impact": impact}

    def logical_delete(self, document_id: str, user_id: str | None = None) -> dict[str, Any]:
        impact = self.preview_delete_impact(document_id, user_id=user_id)
        with db_session(self.db_path) as conn:
            document = get_document(conn, document_id)
            _update(
                conn,
                "documents",
                document_id,
                {"deleted_at": utc_now(), "status": "DELETED"},
            )
            for version in list_versions(conn, document_id):
                deactivate_version_chunks(conn, version["id"])
                _update(
                    conn, "document_versions", version["id"], {"is_active": 0, "is_current": 0}
                )
            add_audit(
                conn,
                event_type="logical_delete",
                actor_user_id=user_id,
                case_id=document["case_id"],
                document_id=document_id,
                decision="allow",
                reason="logical_delete",
                detail_json=json.dumps(impact, ensure_ascii=False),
            )
            return {"status": "DELETED", **impact}

    def resolve_duplicate(
        self, document_id: str, *, action: str, user_id: str | None = None
    ) -> dict[str, Any]:
        with db_session(self.db_path) as conn:
            document = get_document(conn, document_id)
            if not document:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "document not found")
            case_id = document["case_id"]
            current_version_id = document.get("current_version_id")

        self._authorize(user_id, case_id, "material.upload")

        if action == "cancel":
            return self.logical_delete(document_id, user_id=user_id)
        if action == "keep":
            if not current_version_id:
                raise MaterialError(ERROR_CODES["NOT_FOUND"], "no current version")
            with db_session(self.db_path) as conn:
                _update(conn, "documents", document_id, {"status": "UPLOADED"})
            return self.parse_version(current_version_id, user_id=user_id)
        raise MaterialError(ERROR_CODES["PARSE_FAILED"], f"unknown action: {action}")


_default_service: MaterialService | None = None


def get_material_service() -> MaterialService:
    global _default_service
    if _default_service is None:
        _default_service = MaterialService()
    return _default_service


def set_material_service(service: MaterialService | None) -> None:
    global _default_service
    _default_service = service